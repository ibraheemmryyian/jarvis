"""
Document Engine for Jarvis v2
Professional document generation: Word, Excel, PDF, PowerPoint.

Generates formatted business documents, not just text dumps.
"""
import os
from datetime import datetime
from typing import Dict, List, Optional, Any
from .config import WORKSPACE_DIR

# Optional imports with fallbacks
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[DocumentEngine] python-docx not installed. Word generation disabled.")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.chart import BarChart, Reference
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    print("[DocumentEngine] openpyxl not installed. Excel generation disabled.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[DocumentEngine] reportlab not installed. PDF generation disabled.")


class DocumentEngine:
    """
    Professional document generation engine.
    
    Supports:
    - Word (.docx) with styles, headers, tables
    - Excel (.xlsx) with formatting, formulas, charts
    - PDF with professional layouts
    - Templates for common business documents
    """
    
    def __init__(self):
        self.output_dir = os.path.join(WORKSPACE_DIR, "documents")
        os.makedirs(self.output_dir, exist_ok=True)
    
    def get_capabilities(self) -> Dict:
        """Return available document types."""
        return {
            "word": DOCX_AVAILABLE,
            "excel": XLSX_AVAILABLE,
            "pdf": PDF_AVAILABLE,
            "output_dir": self.output_dir
        }
    
    # === WORD DOCUMENTS ===
    
    def create_word_doc(self, filename: str, title: str, 
                        sections: List[Dict]) -> Dict:
        """
        Create a Word document with sections.
        
        Args:
            filename: Output filename (without .docx)
            title: Document title
            sections: List of sections:
                [
                    {"heading": "Introduction", "content": "Text here..."},
                    {"heading": "Analysis", "content": "More text...", "level": 2},
                    {"type": "table", "headers": [...], "rows": [[...], ...]},
                    {"type": "bullet_list", "items": ["Item 1", "Item 2"]}
                ]
        """
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not installed. Run: pip install python-docx"}
        
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Process sections
        for section in sections:
            section_type = section.get("type", "text")
            
            if section_type == "text" or "heading" in section:
                level = section.get("level", 1)
                if section.get("heading"):
                    doc.add_heading(section["heading"], level)
                if section.get("content"):
                    doc.add_paragraph(section["content"])
            
            elif section_type == "table":
                headers = section.get("headers", [])
                rows = section.get("rows", [])
                
                if headers or rows:
                    table_data = [headers] + rows if headers else rows
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            table.rows[i].cells[j].text = str(cell_data)
                            if i == 0:  # Header row
                                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                    
                    doc.add_paragraph()  # Spacer
            
            elif section_type == "bullet_list":
                for item in section.get("items", []):
                    doc.add_paragraph(item, style='List Bullet')
            
            elif section_type == "numbered_list":
                for item in section.get("items", []):
                    doc.add_paragraph(item, style='List Number')
        
        # Save
        filepath = os.path.join(self.output_dir, f"{filename}.docx")
        doc.save(filepath)
        
        return {
            "success": True,
            "path": filepath,
            "format": "docx",
            "sections": len(sections)
        }
    
    def create_report(self, filename: str, title: str, 
                      executive_summary: str, sections: Dict) -> Dict:
        """
        Create a professional report document.
        
        Args:
            filename: Output filename
            title: Report title
            executive_summary: Summary text
            sections: Dict of section_name -> content
        """
        doc_sections = [
            {"heading": "Executive Summary", "content": executive_summary}
        ]
        
        for heading, content in sections.items():
            if isinstance(content, list):
                doc_sections.append({"heading": heading, "type": "bullet_list", "items": content})
            elif isinstance(content, dict) and "headers" in content:
                doc_sections.append({"heading": heading, "type": "table", **content})
            else:
                doc_sections.append({"heading": heading, "content": str(content)})
        
        return self.create_word_doc(filename, title, doc_sections)
    
    # === EXCEL SPREADSHEETS ===
    
    def create_excel(self, filename: str, sheets: List[Dict]) -> Dict:
        """
        Create an Excel workbook with multiple sheets.
        
        Args:
            filename: Output filename (without .xlsx)
            sheets: List of sheet definitions:
                [
                    {
                        "name": "Sheet1",
                        "headers": ["Col A", "Col B", "Col C"],
                        "rows": [[1, 2, 3], [4, 5, 6]],
                        "column_widths": [20, 15, 15]
                    }
                ]
        """
        if not XLSX_AVAILABLE:
            return {"error": "openpyxl not installed. Run: pip install openpyxl"}
        
        wb = Workbook()
        
        # Remove default sheet
        default_sheet = wb.active
        
        for i, sheet_def in enumerate(sheets):
            if i == 0:
                ws = default_sheet
                ws.title = sheet_def.get("name", "Sheet1")
            else:
                ws = wb.create_sheet(sheet_def.get("name", f"Sheet{i+1}"))
            
            # Header style
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            
            # Write headers
            headers = sheet_def.get("headers", [])
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
            
            # Write rows
            rows = sheet_def.get("rows", [])
            for row_idx, row_data in enumerate(rows, 2):
                for col_idx, value in enumerate(row_data, 1):
                    ws.cell(row=row_idx, column=col_idx, value=value)
            
            # Column widths
            widths = sheet_def.get("column_widths", [])
            for col_idx, width in enumerate(widths, 1):
                ws.column_dimensions[chr(64 + col_idx)].width = width
            
            # Add chart if specified
            if sheet_def.get("chart"):
                chart_def = sheet_def["chart"]
                chart = BarChart()
                chart.title = chart_def.get("title", "Chart")
                
                data = Reference(ws, min_col=2, max_col=len(headers),
                               min_row=1, max_row=len(rows) + 1)
                categories = Reference(ws, min_col=1, min_row=2, max_row=len(rows) + 1)
                
                chart.add_data(data, titles_from_data=True)
                chart.set_categories(categories)
                ws.add_chart(chart, "E5")
        
        # Save
        filepath = os.path.join(self.output_dir, f"{filename}.xlsx")
        wb.save(filepath)
        
        return {
            "success": True,
            "path": filepath,
            "format": "xlsx",
            "sheets": len(sheets)
        }
    
    def create_financial_model(self, filename: str, company: str,
                               revenue_data: List, expense_data: List) -> Dict:
        """Create a basic financial model spreadsheet."""
        sheets = [
            {
                "name": "Revenue",
                "headers": ["Month", "Revenue", "Growth %"],
                "rows": revenue_data,
                "column_widths": [15, 15, 12],
                "chart": {"title": "Revenue Growth"}
            },
            {
                "name": "Expenses",
                "headers": ["Category", "Monthly", "Annual"],
                "rows": expense_data,
                "column_widths": [20, 15, 15]
            },
            {
                "name": "Summary",
                "headers": ["Metric", "Value"],
                "rows": [
                    ["Total Revenue", f"=SUM(Revenue!B2:B{len(revenue_data)+1})"],
                    ["Total Expenses", f"=SUM(Expenses!C2:C{len(expense_data)+1})"],
                    ["Net Profit", "=B1-B2"]
                ],
                "column_widths": [20, 20]
            }
        ]
        
        return self.create_excel(filename, sheets)
    
    # === PDF DOCUMENTS ===
    
    def create_pdf(self, filename: str, title: str,
                   content: List[Dict]) -> Dict:
        """
        Create a professional PDF document.
        
        Args:
            filename: Output filename (without .pdf)
            title: Document title
            content: List of content elements:
                [
                    {"type": "heading", "text": "Section 1"},
                    {"type": "paragraph", "text": "Body text..."},
                    {"type": "table", "headers": [...], "rows": [[...]]},
                    {"type": "spacer", "height": 20}
                ]
        """
        if not PDF_AVAILABLE:
            return {"error": "reportlab not installed. Run: pip install reportlab"}
        
        filepath = os.path.join(self.output_dir, f"{filename}.pdf")
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        story = []
        
        # Title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Date
        date_text = f"Generated: {datetime.now().strftime('%B %d, %Y')}"
        story.append(Paragraph(date_text, styles['Normal']))
        story.append(Spacer(1, 24))
        
        # Content
        for item in content:
            item_type = item.get("type", "paragraph")
            
            if item_type == "heading":
                level = item.get("level", 1)
                style = styles[f'Heading{min(level, 3)}']
                story.append(Paragraph(item.get("text", ""), style))
            
            elif item_type == "paragraph":
                story.append(Paragraph(item.get("text", ""), styles['Normal']))
                story.append(Spacer(1, 12))
            
            elif item_type == "table":
                headers = item.get("headers", [])
                rows = item.get("rows", [])
                table_data = [headers] + rows if headers else rows
                
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
                story.append(Spacer(1, 20))
            
            elif item_type == "spacer":
                story.append(Spacer(1, item.get("height", 20)))
            
            elif item_type == "bullet_list":
                for bullet in item.get("items", []):
                    story.append(Paragraph(f"• {bullet}", styles['Normal']))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        
        return {
            "success": True,
            "path": filepath,
            "format": "pdf",
            "elements": len(content)
        }
    
    # === TEMPLATE-BASED GENERATION ===
    
    def create_business_proposal(self, filename: str, company: str,
                                 client: str, proposal: Dict) -> Dict:
        """Generate a formatted business proposal."""
        sections = [
            {"heading": "Executive Summary", "content": proposal.get("summary", "")},
            {"heading": "Problem Statement", "content": proposal.get("problem", "")},
            {"heading": "Proposed Solution", "content": proposal.get("solution", "")},
            {"heading": "Deliverables", "type": "bullet_list", 
             "items": proposal.get("deliverables", [])},
            {"heading": "Timeline", "type": "table",
             "headers": ["Phase", "Duration", "Deliverable"],
             "rows": proposal.get("timeline", [])},
            {"heading": "Investment", "content": proposal.get("pricing", "")},
            {"heading": "Next Steps", "type": "numbered_list",
             "items": proposal.get("next_steps", [])}
        ]
        
        return self.create_word_doc(
            filename,
            f"Business Proposal: {company} for {client}",
            sections
        )
    
    def create_analysis_report(self, filename: str, title: str,
                               analysis: Dict) -> Dict:
        """Generate a formatted analysis report."""
        return self.create_report(
            filename,
            title,
            analysis.get("executive_summary", ""),
            {
                "Key Findings": analysis.get("findings", []),
                "Data Analysis": analysis.get("data", {}),
                "Recommendations": analysis.get("recommendations", []),
                "Appendix": analysis.get("appendix", "")
            }
        )


# Singleton
document_engine = DocumentEngine()
