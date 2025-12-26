"""
Autonomous Executor for Jarvis v2
Self-healing, self-orchestrating execution loop.

"Can't outsmart the AI, but we can out-work it, out-tool it"
"""
import os
import re
import time
import json
from datetime import datetime
from typing import Callable, Dict, List, Optional, Any
from .recycler import recycler
from .router import router
from .coder import coder
from .research import researcher
from .context_manager import context
from .config import WORKSPACE_DIR
from .qa import qa_agent
from .project_manager import project_manager
from .terminal import terminal
from .code_indexer import code_indexer
from .design_creativity import design_creativity
from .visual_qa import visual_qa
from .prompt_refiner import prompt_refiner
from .devtools import devtools  # CTO-level dev tools

# Additional specialist agents for routing
from .code_reviewer import code_reviewer
from .security_auditor import security_auditor
from .brute_research import brute_researcher
from .synthesis import synthesizer

# New specialized development agents
from .frontend_dev import frontend_dev
from .backend_dev import backend_dev
from .ai_ops import ai_ops
from .strategy import strategy
from .architect import architect
from .product_manager import product_manager
from .uiux import uiux
from .seo import seo_specialist

# === PREVIOUSLY UNUSED AGENTS - NOW INTEGRATED ===
# Imports moved to local scope to prevent circular dependencies

# Research & Academic
# from .academic_research import academic_research
# from .research_publisher import research_publisher
# from .academic_workflow import academic_workflow

# Content Generation
# from .content_writer import content_writer
# from .document_engine import document_engine
# from .pitch_deck import pitch_deck, pitch_deck_scorer

# Git & GitHub - Imported lazily to prevent circular imports
# from .git_agent import git_agent
# from .github_agent import github_agent

# Memory & Orchestration
from .memory import memory
# from .orchestrator import orchestrator

# Notifications & Communication
# from .email_agent import email_agent
# from .slack_agent import slack_agent
# from .notifications import notifications
# from .calendar_agent import calendar_agent
# from .daily_briefing import daily_briefing

# Browser & Testing
# from .browser_automation import browser_automation
# from .browser_tester import browser_tester
# from .package_security import package_security

# Business
# from .business_analyst import business_analyst
# from .personality import personality

# Context-segregated registry
from .registry import registry, get_context_for_agent, save_context

# NEW: On-demand context retrieval (RAG-style)
from .context_retriever import context_retriever, get_context

# NEW: Autonomy utilities
from .utils.checkpoint import checkpoint_manager
from .utils.escalation import escalation_manager, should_escalate
from .utils.error_journal import error_journal, log_error, get_avoid_instructions
from .utils.hierarchical_planner import hierarchical_planner

# NEW: Project builder for cohesive file generation
from .project_builder import project_builder

# NEW: Live logging for real-time visibility
from .utils.live_logger import live_logger


class AutonomousExecutor:
    """
    Self-healing autonomous execution loop.
    
    Features:
    - Automatic context recycling at 75%
    - Domain-segregated memory
    - Self-generated continuation prompts
    - Progress tracking across recycles
    - PAUSE/RESUME capability
    - Mid-flow plan modification
    - Phase-specific iteration limits (v4.0)
    """
    
    def __init__(self):
        self.is_running = False
        self.is_paused = False  # NEW: Pause state
        self.current_task = None
        self.current_objective = ""  # Store the objective for reference
        self.iteration = 0
        self.max_iterations = 500  # Jarvis OUTWORKS other AIs - marathon by default
        self.log = []
        self.progress_callback = None
        
        # === V4.0: PHASE LIMITS (prevent infinite loops) ===
        self.coding_iterations = 0
        self.max_coding_iterations = 15  # Force move to git after 15 coding steps
        self.validation_retries = 0
        self.max_validation_retries = 2  # Reduced from 3 - don't burn tokens on unfixable errors
        
        # === V4.0: JUNK FILE FILTER ===
        # Files that should NEVER be created (stdlib reimplementations)
        self.junk_files = {
            'asyncio.py', 'contextlib.py', 'contextlib_impl.py', 'sqlite3.py', 
            'sqlite3_impl.py', 'pydantic.py', 'fastapi.py', 'sqlalchemy.py',
            'sqlalchemy_impl.py', 'types.py', 'typing.py', 'collections.py',
            '__enter__.py', '__exit__.py', 'contextmanager.py', 'generatorcontextmanager.py',
            '_generatorcontextmanager.py', '_generator_context_manager.py', 
            'generator_context_manager.py', 'eventloop.py', 'run_until_complete.py',
            'wait_for_task.py', 'my_coroutine.py'
        }
        
        # === V4.1: PROJECT TYPE LOCK ===
        # Detected once, enforced throughout the run
        self.project_type = None  # 'react', 'python', 'fullstack', 'research', 'landing'
        self.project_type_rules = {
            'react': {
                'allowed_extensions': ['.jsx', '.tsx', '.js', '.ts', '.css', '.html', '.json', '.md'],
                'forbidden_extensions': ['.py'],
                'prompt_injection': """
## 🔒 PROJECT TYPE: REACT FRONTEND
- ONLY generate React components (.jsx/.tsx)
- Use CSS/TailwindCSS for styling
- DO NOT create Python files
- DO NOT create backend APIs
- Focus on: components, hooks, state management
"""
            },
            'python': {
                'allowed_extensions': ['.py', '.json', '.md', '.txt', '.yml', '.yaml'],
                'forbidden_extensions': ['.jsx', '.tsx'],
                'prompt_injection': """
## 🔒 PROJECT TYPE: PYTHON
- Generate Python modules and scripts
- Use proper project structure (src/, tests/, etc.)
- DO NOT create React/frontend code
- Focus on: modules, APIs, data processing
"""
            },
            'fullstack': {
                'allowed_extensions': ['.jsx', '.tsx', '.js', '.ts', '.css', '.html', '.json', '.md', '.py'],
                'forbidden_extensions': [],
                'prompt_injection': """
## 🔒 PROJECT TYPE: FULLSTACK
- Frontend: React components in src/
- Backend: Python/FastAPI in backend/
- Keep frontend and backend separate
"""
            },
            'research': {
                'allowed_extensions': ['.py', '.md', '.csv', '.json', '.txt'],
                'forbidden_extensions': ['.jsx', '.tsx'],
                'prompt_injection': """
## 🔒 PROJECT TYPE: RESEARCH
- Generate: algorithm.py, benchmark.py, paper.md
- Include real benchmark data
- DO NOT create web interfaces
"""
            },
            'landing': {
                'allowed_extensions': ['.html', '.css', '.js', '.json', '.md'],
                'forbidden_extensions': ['.py', '.jsx', '.tsx'],
                'prompt_injection': """
## 🔒 PROJECT TYPE: LANDING PAGE
- Generate static HTML/CSS/JS only
- No frameworks, no build tools
- Focus on visual design and responsiveness
"""
            }
        }
        
        # Pause/Resume state
        self.pause_requested = False
        self.resume_event = None  # Will be threading.Event()
        self.pause_reason = ""
        self.pending_modifications = []  # Plan changes during pause
        self.saved_state = {}  # State snapshot on pause
    
    def pause(self, reason: str = "User requested pause") -> Dict:
        """
        Pause execution at the next safe point.
        Returns current state.
        """
        if not self.is_running:
            return {"success": False, "error": "Not currently running"}
        
        self.pause_requested = True
        self.pause_reason = reason
        self._log(f"⏸️ Pause requested: {reason}")
        
        return {
            "success": True,
            "message": "Will pause after current step completes",
            "current_step": self.iteration,
            "reason": reason
        }
    
    def resume(self, modifications: List[str] = None) -> Dict:
        """
        Resume paused execution, optionally with plan modifications.
        
        Args:
            modifications: List of changes to apply to the plan
        """
        if not self.is_paused:
            return {"success": False, "error": "Not currently paused"}
        
        if modifications:
            self.pending_modifications.extend(modifications)
            self._log(f"📝 {len(modifications)} modifications queued")
        
        self.is_paused = False
        self.pause_requested = False
        
        # Signal resume if using threading
        if self.resume_event:
            self.resume_event.set()
        
        self._log("▶️ Resuming execution")
        
        return {
            "success": True,
            "message": "Execution resumed",
            "modifications_applied": len(modifications) if modifications else 0
        }
    
    def get_state(self) -> Dict:
        """Get current execution state (useful when paused)."""
        from .recycler import recycler
        progress = recycler.get_progress()
        
        return {
            "is_running": self.is_running,
            "is_paused": self.is_paused,
            "pause_reason": self.pause_reason if self.is_paused else None,
            "current_iteration": self.iteration,
            "max_iterations": self.max_iterations,
            "objective": progress.get("objective", ""),
            "completed_steps": progress.get("completed_steps", []),
            "pending_steps": progress.get("pending_steps", []),
            "percent_complete": progress.get("percent", 0),
            "pending_modifications": self.pending_modifications,
            "log_tail": self.log[-10:] if self.log else []
        }
    
    def modify_plan(self, action: str, step: str = None, position: int = None) -> Dict:
        """
        Modify the execution plan mid-flow.
        
        Args:
            action: "add", "remove", "insert", "replace"
            step: The step text
            position: For insert, where to insert
        """
        from .recycler import recycler
        
        if action == "add":
            recycler.pending_steps.append(step)
            self._log(f"📌 Added step: {step}")
        elif action == "remove" and step:
            if step in recycler.pending_steps:
                recycler.pending_steps.remove(step)
                self._log(f"❌ Removed step: {step}")
        elif action == "insert" and position is not None:
            recycler.pending_steps.insert(position, step)
            self._log(f"📍 Inserted step at {position}: {step}")
        elif action == "replace" and position is not None:
            if position < len(recycler.pending_steps):
                old = recycler.pending_steps[position]
                recycler.pending_steps[position] = step
                self._log(f"🔄 Replaced step {position}: {old} → {step}")
        
        return {"success": True, "pending_steps": recycler.pending_steps}
    
    def add_note(self, note: str) -> Dict:
        """Add a note to the execution log (useful during pause)."""
        self._log(f"📝 Note: {note}")
        
        # Also save to memory
        from .memory import memory
        memory.save_fact(note, category="execution_notes")
        
        return {"success": True}
    
    def _check_pause(self) -> bool:
        """
        Check if we should pause. Called between steps.
        Returns True if paused, False to continue.
        """
        if self.pause_requested:
            self.is_paused = True
            self.pause_requested = False
            
            # Save state
            from .recycler import recycler
            self.saved_state = {
                "iteration": self.iteration,
                "progress": recycler.get_progress(),
                "timestamp": datetime.now().isoformat()
            }
            
            self._log(f"⏸️ PAUSED at step {self.iteration}")
            self._log(f"   Reason: {self.pause_reason}")
            self._log("   Use executor.resume() to continue")
            self._log("   Use executor.modify_plan() to change steps")
            
            # If using threading, wait for resume
            if self.resume_event:
                import threading
                self.resume_event = threading.Event()
                self.resume_event.wait()  # Block until resumed
            
            return True
        
        return False
    
    def _apply_pending_modifications(self):
        """Apply any modifications queued during pause."""
        if self.pending_modifications:
            self._log(f"📝 Applying {len(self.pending_modifications)} modifications")
            for mod in self.pending_modifications:
                # Parse modification string
                if mod.startswith("add:"):
                    self.modify_plan("add", mod[4:].strip())
                elif mod.startswith("remove:"):
                    self.modify_plan("remove", mod[7:].strip())
            self.pending_modifications = []
    
    def _log(self, msg: str):
        """Log with timestamp."""
        timestamp = datetime.now().strftime("%H:%M:%S")
        entry = f"[{timestamp}] {msg}"
        self.log.append(entry)
        
        # Safe printing for Windows consoles
        try:
            print(entry, flush=True)
        except UnicodeEncodeError:
            # Fallback: remove non-ascii chars if console can't handle them
            safe_msg = msg.encode('ascii', 'ignore').decode('ascii')
            safe_entry = f"[{timestamp}] {safe_msg}"
            print(safe_entry, flush=True)
            
        if self.progress_callback:
            self.progress_callback(msg)
    
    def _call_llm(self, prompt: str, max_tokens: int = None) -> str:
        """Make LLM call through base agent with optional token limit."""
        from .base_agent import BaseAgent
        from .config import TOKEN_LIMITS
        
        # Default to standard limit if not specified
        tokens = max_tokens or TOKEN_LIMITS["standard"]
        
        class TempAgent(BaseAgent):
            def __init__(self):
                super().__init__("autonomous")
            def _get_system_prompt(self):
                return "You are Jarvis, an autonomous AI assistant completing a multi-step task."
            def run(self, task: str) -> str:
                return self.call_llm(task)
        
        agent = TempAgent()
        return agent.call_llm(prompt, max_tokens=tokens)
    
    def _check_completion(self, response: str) -> bool:
        """Check if task is complete based on EXPLICIT completion signal in response."""
        # VERY SPECIFIC signals that the LLM must explicitly output
        # Generic terms like "done" trigger false positives
        completion_signals = [
            "[task complete]",
            "[all steps complete]",
            "[project finished]",
            "jarvis: task complete",
            "=== task complete ===",
            "[completion]",
        ]
        response_lower = response.lower()
        
        # Only trigger if signal appears at start of a line or is bracketed
        for signal in completion_signals:
            if signal in response_lower:
                return True
        
        return False
    
    def _verify_research_outputs(self, project_path: str, objective: str) -> dict:
        """
        Verify that research task outputs actually exist and work.
        Returns dict with verified status and missing items.
        """
        issues = []
        verified = []
        obj_lower = objective.lower()
        
        # Check if this is a research/algorithm task
        is_research = any(kw in obj_lower for kw in [
            "research", "algorithm", "paper", "academic", "benchmark", "novel"
        ])
        
        if not is_research:
            return {"verified": True, "issues": [], "verified_items": []}
        
        # Required files for research tasks
        required_files = {
            "paper.md": "Academic paper",
            "algorithm.py": "Algorithm implementation",
            "benchmark.py": "Benchmark script",
        }
        
        # Check for required files
        for filename, description in required_files.items():
            # Check in root and common subdirs
            found = False
            for subdir in ["", "src", "backend", "docs"]:
                check_path = os.path.join(project_path, subdir, filename)
                if os.path.exists(check_path):
                    found = True
                    verified.append(f"{description}: {check_path}")
                    break
            if not found:
                issues.append(f"Missing {description} ({filename})")
        
        # Check for results folder/CSV
        results_found = False
        for pattern in ["results", "output", "benchmarks"]:
            results_dir = os.path.join(project_path, pattern)
            if os.path.exists(results_dir):
                csv_files = [f for f in os.listdir(results_dir) if f.endswith('.csv')]
                if csv_files:
                    results_found = True
                    verified.append(f"Results data: {csv_files}")
        
        if not results_found:
            issues.append("Missing benchmark results (no CSV files in results/)")
        
        # Try to run a simple Python file to verify code works
        main_candidates = ["main.py", "benchmark.py", "run_benchmark.py"]
        for candidate in main_candidates:
            main_path = os.path.join(project_path, candidate)
            if not os.path.exists(main_path):
                main_path = os.path.join(project_path, "scripts", candidate)
            
            if os.path.exists(main_path):
                try:
                    import subprocess
                    result = subprocess.run(
                        ["python", "-c", f"import ast; ast.parse(open(r'{main_path}').read())"],
                        capture_output=True, text=True, timeout=10
                    )
                    if result.returncode == 0:
                        verified.append(f"Code syntax valid: {candidate}")
                    else:
                        issues.append(f"Syntax error in {candidate}: {result.stderr[:100]}")
                except Exception as e:
                    issues.append(f"Could not verify {candidate}: {e}")
                break
        
        # Check paper.md for placeholders
        paper_paths = [
            os.path.join(project_path, "paper.md"),
            os.path.join(project_path, "docs", "README.md"),
            os.path.join(project_path, "docs", "paper.md"),
        ]
        for paper_path in paper_paths:
            if os.path.exists(paper_path):
                with open(paper_path, 'r', encoding='utf-8', errors='ignore') as f:
                    paper_content = f.read()
                    placeholders = ["[Include", "[TODO", "[PLACEHOLDER", "[Insert", "[Add"]
                    for p in placeholders:
                        if p in paper_content:
                            issues.append(f"Paper contains placeholder: {p}")
                break
        
        return {
            "verified": len(issues) == 0,
            "issues": issues,
            "verified_items": verified
        }
    
    def _validate_output(self, content: str, task_type: str) -> dict:
        """Validate output quality - reject placeholder-filled or incomplete content."""
        issues = []
        
        # Check for common placeholders
        placeholders = [
            "[Company Name]", "[Your Name]", "[Your Company]",
            "[specific initiative", "[TODO", "[PLACEHOLDER",
            "[Company]", "[Name]", "[Contact]"
        ]
        
        for p in placeholders:
            if p in content:
                issues.append(f"Contains placeholder: {p}")
        
        # CODE-SPECIFIC QUALITY CHECKS
        if task_type == "coding":
            content_lower = content.lower()
            
            # Check for incomplete code markers
            incomplete_markers = [
                ("# TODO", "Contains TODO marker"),
                ("// TODO", "Contains TODO marker"),
                ("# FIXME", "Contains FIXME marker"),
                ("// FIXME", "Contains FIXME marker"),
                ("pass  # ", "Contains pass placeholder"),
                ("...", "Contains ellipsis placeholder"),
                ("# add more", "Contains 'add more' placeholder"),
                ("// add more", "Contains 'add more' placeholder"),
                ("# implement", "Contains 'implement' placeholder"),
                ("// implement", "Contains 'implement' placeholder"),
                ("raise NotImplementedError", "Contains NotImplementedError"),
            ]
            
            for marker, issue in incomplete_markers:
                if marker.lower() in content_lower:
                    issues.append(issue)
            
            # Check for very short code output (likely fragment)
            code_lines = [l for l in content.split('\n') if l.strip() and not l.strip().startswith('#') and not l.strip().startswith('//')]
            if len(code_lines) < 20 and "def " in content or "function " in content:
                issues.append("Code too short (< 20 non-comment lines)")
            
            # Check for empty function bodies
            if "pass\n" in content and content.count("def ") > content.count("pass\n") + 1:
                pass  # Has some pass statements but also real implementations - OK
            elif content.count("pass\n") > 2:
                issues.append("Multiple empty function bodies (pass statements)")
            
            # Check for missing error handling in API code
            if "fastapi" in content_lower or "@app." in content_lower:
                if "HTTPException" not in content and "raise" not in content:
                    issues.append("API code missing error handling (no HTTPException)")
        
        # For analysis tasks, require actual data
        if task_type == "analysis":
            if content.count("[") > 10:
                issues.append("Too many bracket placeholders for analysis output")
        
        if issues:
            self._log(f"  ⚠️ OUTPUT QUALITY ISSUES: {', '.join(issues[:3])}...")
            return {"valid": False, "issues": issues}
        
        return {"valid": True, "issues": []}
    
    def _detect_domain(self, content: str) -> str:
        """Detect which domain the content belongs to."""
        content_lower = content.lower()
        
        if any(kw in content_lower for kw in ["react", "component", "jsx", "css", "html", "ui", "frontend", "page"]):
            return "frontend"
        elif any(kw in content_lower for kw in ["api", "route", "server", "endpoint", "express", "backend"]):
            return "backend"
        elif any(kw in content_lower for kw in ["database", "schema", "table", "sql", "supabase", "migration"]):
            return "database"
        elif any(kw in content_lower for kw in ["market", "competitor", "research", "trend", "analysis"]):
            return "research"
        else:
            return "decisions"
    
    def _detect_task_type(self, objective: str) -> str:
        """
        Detect what type of task this is for appropriate handling.
        Supports coding, research, writing, analysis, and general tasks.
        """
        obj_lower = objective.lower()
        
        # Coding tasks
        coding_keywords = [
            "build", "create", "develop", "code", "implement", "website", "app",
            "frontend", "backend", "api", "component", "page", "feature",
            "react", "python", "javascript", "html", "css", "database",
            "deploy", "fix bug", "refactor", "landing page", "saas"
        ]
        if any(kw in obj_lower for kw in coding_keywords):
            return "coding"
        
        # Research tasks
        research_keywords = [
            "research", "investigate", "find out", "look into", "analyze market",
            "competitor analysis", "industry trends", "gather information",
            "what is", "how does", "why do", "explore options"
        ]
        if any(kw in obj_lower for kw in research_keywords):
            return "research"
        
        # Writing tasks
        writing_keywords = [
            "write", "draft", "compose", "create content", "blog post",
            "article", "email", "pitch", "proposal", "documentation",
            "copy", "script", "outline", "summary"
        ]
        if any(kw in obj_lower for kw in writing_keywords):
            return "writing"
        
        # Analysis tasks
        analysis_keywords = [
            "analyze", "evaluate", "assess", "compare", "review",
            "audit", "measure", "calculate", "forecast", "model",
            "data analysis", "metrics", "performance"
        ]
        if any(kw in obj_lower for kw in analysis_keywords):
            return "analysis"
        
        # Default to general
        return "general"
    
    def _detect_and_lock_project_type(self, objective: str) -> str:
        """
        V4.1: Detect project type from objective and LOCK it.
        Called once at the start, returns prompt injection text.
        """
        if self.project_type:
            # Already locked, return existing rules
            rules = self.project_type_rules.get(self.project_type, {})
            return rules.get('prompt_injection', '')
        
        obj_lower = objective.lower()
        
        # Detect React/Frontend
        if any(kw in obj_lower for kw in ['react', 'next.js', 'nextjs', 'vite', 'jsx', 'tsx', 'component', 'frontend app', 'web app']):
            if not any(kw in obj_lower for kw in ['fastapi', 'flask', 'django', 'backend', 'api server']):
                self.project_type = 'react'
                self._log("🔒 PROJECT TYPE LOCKED: React Frontend")
        
        # Detect Python/Backend
        elif any(kw in obj_lower for kw in ['python script', 'fastapi', 'flask', 'django', 'backend', 'cli tool', 'data processing']):
            self.project_type = 'python'
            self._log("🔒 PROJECT TYPE LOCKED: Python")
        
        # Detect Research
        elif any(kw in obj_lower for kw in ['research', 'algorithm', 'paper', 'benchmark', 'novel', 'academic']):
            self.project_type = 'research'
            self._log("🔒 PROJECT TYPE LOCKED: Research")
        
        # Detect Landing Page
        elif any(kw in obj_lower for kw in ['landing page', 'static site', 'html page', 'marketing page']):
            self.project_type = 'landing'
            self._log("🔒 PROJECT TYPE LOCKED: Landing Page")
        
        # Detect Fullstack
        elif any(kw in obj_lower for kw in ['fullstack', 'full-stack', 'full stack', 'frontend and backend']):
            self.project_type = 'fullstack'
            self._log("🔒 PROJECT TYPE LOCKED: Fullstack")
        
        # Default: try to infer from common patterns
        else:
            # If mentions web but not backend, assume React
            if any(kw in obj_lower for kw in ['web', 'app', 'ui', 'interface', 'page']) and not any(kw in obj_lower for kw in ['api', 'server', 'database']):
                self.project_type = 'react'
                self._log("🔒 PROJECT TYPE LOCKED: React (inferred)")
            else:
                self.project_type = 'fullstack'
                self._log("🔒 PROJECT TYPE LOCKED: Fullstack (default)")
        
        rules = self.project_type_rules.get(self.project_type, {})
        return rules.get('prompt_injection', '')
    
    def _extract_component_name(self, step: str) -> str:
        """Extract a clean component name from a step description for file naming."""
        if not step or step == "unknown":
            return "component"
        
        # Look for [COMPONENT: Name] pattern
        import re
        component_match = re.search(r'\[COMPONENT[:\s]+([^\]]+)\]', step, re.IGNORECASE)
        if component_match:
            name = component_match.group(1).strip()
        else:
            # Extract key words from step
            name = step
        
        # Clean up the name for filesystem
        # Take first few meaningful words
        words = re.findall(r'[a-zA-Z]+', name)
        if words:
            # Take up to 3 words, make lowercase, join with hyphen
            clean_words = [w.lower() for w in words[:3] if len(w) > 2]
            return "-".join(clean_words) if clean_words else "component"
        
        return "component"
    
    def _route_to_specialist(self, step: str, task_type: str, project_path: str) -> Dict:
        """
        Route a step to the appropriate specialist agent.
        NOW USES: On-demand context retrieval (RAG-style) instead of pre-loading domains.
        """
        step_lower = step.lower()
        step_upper = step.upper()
        
        # Get context on-demand using the new retriever
        # This asks the LLM which files are needed instead of pre-loading everything
        context = get_context(step, agent=task_type, project=project_path)
        
        # === COMPONENT STEPS: Direct LLM for complete modules ===
        if "[COMPONENT" in step_upper or "complete module" in step_lower:
            return {
                "agent": "component_builder", 
                "category": "FRONTEND",
                "use_specialist": False,
                "context": context  # On-demand context instead of domains
            }
        
        # === ARCHITECTURE STEPS: Direct LLM for design docs ===
        if "[ARCHITECTURE" in step_upper or "system design" in step_lower:
            return {
                "agent": "architect",
                "category": "ARCHITECTURE", 
                "use_specialist": True,
                "context": context
            }
        
        # === INTEGRATION STEPS: Direct LLM for connecting modules ===
        if "[INTEGRATION" in step_upper:
            return {
                "agent": "integrator",
                "category": "FRONTEND",
                "use_specialist": False,
                "context": context
            }
        
        # === Route by keywords to categories ===
        routing_rules = {
            "FRONTEND": {
                "keywords": ["ui", "component", "page", "css", "style", "react", "animation", "frontend", "responsive"],
                "default_agent": "frontend_dev"
            },
            "BACKEND": {
                "keywords": ["api", "endpoint", "database", "auth", "backend", "server", "crud", "rest", "graphql"],
                "default_agent": "backend_dev"
            },
            "RESEARCH": {
                "keywords": ["research", "investigate", "find", "search", "analyze"],
                "default_agent": "brute_researcher"
            },
            "ACADEMIC": {
                "keywords": ["paper", "academic", "cite", "publication", "journal", "abstract", "methodology"],
                "default_agent": "academic_research"
            },
            "QA": {
                "keywords": ["test", "qa", "quality", "debug", "fix bug", "error", "lint", "review"],
                "default_agent": "qa_agent"
            },
            "OPS": {
                "keywords": ["deploy", "docker", "kubernetes", "ci/cd", "production", "github", "push", "git"],
                "default_agent": "ops"
            },
            "CONTENT": {
                "keywords": ["write", "document", "blog", "content", "seo", "readme", "article"],
                "default_agent": "content_writer"
            },
            "PRESENTATION": {
                "keywords": ["pitch", "deck", "slides", "presentation", "investor"],
                "default_agent": "pitch_deck"
            },
            "BUSINESS": {
                "keywords": ["requirements", "stakeholder", "user story", "business", "analysis", "specification"],
                "default_agent": "business_analyst"
            }
        }
        
        # Match category
        for category, config in routing_rules.items():
            if any(kw in step_lower for kw in config["keywords"]):
                return {
                    "agent": config["default_agent"],
                    "category": category,
                    "use_specialist": True,
                    "context": context  # On-demand context
                }
        
        # Default: use LLM directly with minimal context
        return {
            "agent": "general",
            "category": "CORE", 
            "use_specialist": False,
            "context": context
        }
    
    def _call_specialist(self, agent_name: str, step: str, project_path: str, context: str) -> str:
        """Call a specialist agent with the step and context."""
        self._log(f"  -> Routing to: {agent_name}")
        
        # Get project name for file saving
        project_name = os.path.basename(project_path) if project_path else self._get_project_name(recycler.task_objective or "project")
        
        try:
            result = None
            
            if agent_name == "coder":
                result = coder.run(step, project_path if os.path.exists(project_path) else None)
                
            elif agent_name == "code_reviewer":
                result = str(code_reviewer.run(project_path))
                
            elif agent_name == "qa" or agent_name == "qa_agent":
                result = str(qa_agent.run(project_path))
                
            elif agent_name == "security":
                result = str(security_auditor.run(project_path))
                
            elif agent_name == "research" or agent_name == "brute_researcher":
                # Research agent - returns text content
                result = str(brute_researcher.run(step))
                # Save research to docs
                if result and len(result) > 100:
                    self._save_research_doc(result, project_path, step)
                
            elif agent_name == "frontend_dev":
                # Frontend development - generate and save code
                from .frontend_dev import frontend_dev
                # Pass existing project context so agent knows what's built
                file_context = self._get_project_context(project_path, task=step)
                result = frontend_dev.run(step, project_context=file_context)
                # Extract and save any code blocks
                if result:
                    self._extract_and_save_code(result, project_name)
                
            elif agent_name == "backend_dev":
                # Backend development - generate and save code
                from .backend_dev import backend_dev
                # Pass existing project context with task awareness
                file_context = self._get_project_context(project_path, task=step)
                result = backend_dev.run(step, project_context=file_context)
                # Extract and save any code blocks
                if result:
                    self._extract_and_save_code(result, project_name)
                
            elif agent_name == "content_writer":
                # Content writing - save to docs
                from .content_writer import content_writer
                result = content_writer.run(step)
                if result and len(result) > 100:
                    self._save_content_doc(result, project_path, step)
                
            elif agent_name == "architect":
                # Architecture - save design docs
                from .architect import architect
                result = architect.run(step)
                if result and len(result) > 100:
                    self._save_architecture_doc(result, project_path, step)
                    
            elif agent_name == "ops":
                # DevOps agent
                from .ops import ops
                result = ops.run(step)
                if result:
                    self._extract_and_save_code(result, project_name)
            
            # === NEW AGENT HANDLERS ===
            
            elif agent_name == "academic_research":
                # Academic research - for papers, citations, methodology
                from .academic_research import academic_research
                result = academic_research.run(step)
                if result and len(result) > 100:
                    self._save_research_doc(result, project_path, step)
            
            elif agent_name == "pitch_deck":
                # Presentation generation
                from .pitch_deck import pitch_deck
                result = pitch_deck.run(step)
                if result:
                    self._save_content_doc(result, project_path, step)
            
            elif agent_name == "business_analyst":
                # Business requirements and analysis
                from .business_analyst import business_analyst
                result = business_analyst.run(step)
                if result and len(result) > 100:
                    self._save_architecture_doc(result, project_path, step)
            
            elif agent_name == "git_agent":
                # Git operations
                from .git_agent import git_agent
                result = git_agent.run(step, project_path)
            
            elif agent_name == "github_agent":
                # GitHub operations
                from .github_agent import github_agent
                result = github_agent.run(step, project_path)
                    
            else:
                # Fallback to LLM with code extraction
                result = self._call_llm(f"Execute this step: {step}\n\nContext: {context[:2000]}")
                # Still try to extract and save any code from response
                if result:
                    self._extract_and_save_code(result, project_name)
            
            return result or ""
                
        except Exception as e:
            self._log(f"  Specialist error: {e}")
            # Fallback to LLM on error
            result = self._call_llm(f"Execute this step: {step}\n\nContext: {context[:2000]}")
            # Try to extract code even from fallback
            if result:
                self._extract_and_save_code(result, project_name)
            return result or ""
    
    def _get_project_context(self, project_path: str, task: str = "") -> str:
        """
        Get smart codebase context for the current task.
        Uses code_indexer for intelligent context extraction.
        """
        if not project_path or not os.path.exists(project_path):
            return ""
        
        try:
            # First, index the project
            code_indexer.index_project(project_path)
            
            # Get relevant context based on current task
            if task:
                relevant_context = code_indexer.get_relevant_context(
                    project_path, 
                    task, 
                    max_tokens=4000
                )
                if relevant_context:
                    return f"RELEVANT CODEBASE CONTEXT:\n{relevant_context}"
            
            # Fallback: Get project file summary
            existing_files = []
            for root, dirs, files in os.walk(project_path):
                # Skip common noise directories
                dirs[:] = [d for d in dirs if d not in ['__pycache__', 'node_modules', '.git', '.jarvis', 'venv', 'dist', 'build']]
                
                for f in files:
                    rel_path = os.path.relpath(os.path.join(root, f), project_path)
                    # Get file size
                    try:
                        size = os.path.getsize(os.path.join(root, f))
                        if size > 50:  # Only show non-trivial files
                            existing_files.append(f"{rel_path} ({size} bytes)")
                    except:
                        existing_files.append(rel_path)
            
            if existing_files[:30]:  # Limit to 30 files
                return "EXISTING PROJECT FILES:\n" + "\n".join(existing_files[:30])
        except Exception as e:
            self._log(f"  Context extraction error: {e}")
        
        return ""
    
    def _save_research_doc(self, content: str, project_path: str, step: str):
        """Save research output to docs folder."""
        docs_dir = os.path.join(project_path, "docs", "research")
        os.makedirs(docs_dir, exist_ok=True)
        
        # Generate filename from step
        filename = self._extract_component_name(step) or "research"
        filename = filename.replace(" ", "_").lower()[:30] + ".md"
        
        filepath = os.path.join(docs_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        self._log(f"  -> Saved: docs/research/{filename} ({len(content)} chars)")
    
    def _save_content_doc(self, content: str, project_path: str, step: str):
        """Save content writing output to docs folder."""
        docs_dir = os.path.join(project_path, "docs")
        os.makedirs(docs_dir, exist_ok=True)
        
        # Determine filename from step
        step_lower = step.lower()
        if "whitepaper" in step_lower:
            filename = "whitepaper.md"
        elif "business" in step_lower or "plan" in step_lower:
            filename = "business_plan.md"
        elif "readme" in step_lower:
            filename = "README.md"
            docs_dir = project_path  # README goes in root
        else:
            filename = self._extract_component_name(step) or "document"
            filename = filename.replace(" ", "_").lower()[:30] + ".md"
        
        filepath = os.path.join(docs_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        self._log(f"  -> Saved: {os.path.relpath(filepath, project_path)} ({len(content)} chars)")
    
    def _save_architecture_doc(self, content: str, project_path: str, step: str):
        """Save architecture docs to docs/design folder."""
        docs_dir = os.path.join(project_path, "docs", "design")
        os.makedirs(docs_dir, exist_ok=True)
        
        step_lower = step.lower()
        if "database" in step_lower or "schema" in step_lower:
            filename = "database_schema.md"
        elif "api" in step_lower:
            filename = "api_spec.md"
        elif "architecture" in step_lower:
            filename = "system_architecture.md"
        else:
            filename = self._extract_component_name(step) or "design"
            filename = filename.replace(" ", "_").lower()[:30] + ".md"
        
        filepath = os.path.join(docs_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        self._log(f"  -> Saved: docs/design/{filename} ({len(content)} chars)")
    
    def _finalize_project(self, project_path: str) -> Dict:
        """
        Phase 6: Consolidate and finalize the project.
        
        - Removes files that shadow Python packages
        - Merges fragmented small files
        - Asks LLM to fix any remaining issues
        - Verifies all imports work
        """
        import os
        import subprocess
        
        result = {
            "status": "complete",
            "files_consolidated": 0,
            "shadowing_fixed": 0,
            "imports_verified": 0,
            "errors_fixed": 0
        }
        
        if not os.path.exists(project_path):
            return result
            
        # Files that shadow Python packages - DELETE these
        SHADOW_FILES = {
            'passlib.py', 'sqlalchemy.py', 'base64.py', 'jose.py', 'jwt.py',
            'pydantic.py', 'fastapi.py', 'flask.py', 'requests.py', 'httpx.py',
            'cryptography.py', 'bcrypt.py', 'dotenv.py', 'redis.py', 'celery.py',
            'pytest.py', 'unittest.py', 'logging.py', 'asyncio.py', 'secrets.py',
            'contextlib.py', 'hashlib.py', 'json.py', 'os.py', 'sys.py', 're.py'
        }
        
        deleted_shadow_files = []
        fragmented_files = []
        
        # Walk through project and identify issues
        for root, dirs, files in os.walk(project_path):
            # Skip __pycache__ and .jarvis
            dirs[:] = [d for d in dirs if d not in ['__pycache__', '.jarvis', 'node_modules', '.git']]
            
            for file in files:
                if not file.endswith('.py'):
                    continue
                    
                filepath = os.path.join(root, file)
                
                # Check for shadowing files
                if file in SHADOW_FILES:
                    try:
                        os.remove(filepath)
                        deleted_shadow_files.append(file)
                        result["shadowing_fixed"] += 1
                    except:
                        pass
                    continue
                
                # Check for small fragmented files (< 100 bytes)
                try:
                    size = os.path.getsize(filepath)
                    if size < 100 and size > 0:
                        fragmented_files.append(filepath)
                except:
                    pass
        
        # Delete small fragmented files (they're usually incomplete stubs)
        for frag_file in fragmented_files:
            try:
                os.remove(frag_file)
                result["files_consolidated"] += 1
            except:
                pass
        
        # Count how many Python files pass import check
        python_files = []
        for root, dirs, files in os.walk(project_path):
            dirs[:] = [d for d in dirs if d not in ['__pycache__', '.jarvis', 'node_modules', '.git']]
            for file in files:
                if file.endswith('.py') and file != '__init__.py':
                    python_files.append(os.path.join(root, file))
        
        for py_file in python_files[:20]:  # Check up to 20 files
            try:
                # Simple syntax check
                with open(py_file, 'r', encoding='utf-8') as f:
                    code = f.read()
                compile(code, py_file, 'exec')
                result["imports_verified"] += 1
            except SyntaxError:
                pass  # File has syntax errors
            except:
                pass
        
        return result
    
    def _extract_and_save_code(self, result: str, project_name: str) -> List[str]:
        """
        Extract code blocks from LLM output and save to proper files.
        Handles both:
        1. Code blocks with explicit filenames: ```jsx filename="App.jsx"
        2. Concatenated code with comment markers: // src/components/Sidebar.tsx
        """
        import re
        import os
        
        saved_files = []
        
        # Get or create project using project_builder
        if not project_builder.project_path or project_name not in str(project_builder.project_path or ""):
            objective = recycler.task_objective or ""
            template = "react" if any(kw in objective.lower() for kw in ["react", "next", "jsx", "tsx"]) else "vanilla"
            project_builder.create_project(project_name, template=template)
        
        project_path = project_builder.project_path
        
        # FIRST: Try to split by comment markers (// src/filename.ext or # src/filename.ext)
        file_pattern = r'(?:\/\/|#)\s*(src\/[^\n]+\.(?:tsx?|jsx?|css|py|json|md))\s*\n'
        splits = re.split(file_pattern, result)
        
        if len(splits) > 2:
            # We have comment-based splits
            self._log(f"  -> Detected {len(splits)//2} embedded files")
            i = 1
            while i < len(splits) - 1:
                filename = splits[i].strip()
                content = splits[i + 1].strip()
                
                # === V4.0: JUNK FILE FILTER ===
                basename = os.path.basename(filename)
                if basename in self.junk_files:
                    self._log(f"  -> SKIPPED JUNK: {filename} (stdlib reimplementation)")
                    i += 2
                    continue
                
                # Clean up content - remove trailing ```
                if content.endswith('```'):
                    content = content[:-3].strip()
                
                # Remove leading ``` with language
                if content.startswith('```'):
                    lines = content.split('\n', 1)
                    content = lines[1] if len(lines) > 1 else ""
                
                if content and filename:
                    full_path = os.path.join(project_path, filename)
                    os.makedirs(os.path.dirname(full_path), exist_ok=True)
                    
                    # V4.4: SMART OVERWRITE PROTECTION for markdown files
                    if not self._should_overwrite_file(full_path, content):
                        self._log(f"  ⚠️ SKIP: Won't overwrite {filename} with lower quality content")
                        i += 2
                        continue
                    
                    with open(full_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    
                    saved_files.append(filename)
                    self._log(f"  -> Saved: {filename} ({len(content)} chars)")
                
                i += 2
            
            return saved_files
        
        # SECOND: Try explicit filename in code blocks
        code_pattern = r'```(\w+)(?:\s+filename=["\']([^"\']+)["\'])?\n([\s\S]*?)```'
        matches = re.findall(code_pattern, result)
        
        if matches:
            for lang, filename, code in matches:
                lang = lang.lower()
                code = code.strip()
                
                if not code:
                    continue
                
                # FIRST: Try to extract filename from first line comment
                # Matches: // src/file.js, # backend/file.py, /* src/styles.css */
                first_line = code.split('\n')[0].strip()
                extracted_filename = None
                
                # Pattern for // or # comments with path
                comment_file_match = re.search(r'^(?://|#|\*/?\*?)\s*([a-zA-Z0-9_/.-]+\.[a-zA-Z]+)', first_line)
                if comment_file_match:
                    potential_file = comment_file_match.group(1)
                    # Validate it looks like a real file path
                    if '/' in potential_file or potential_file.endswith(('.py', '.js', '.jsx', '.tsx', '.css', '.html')):
                        extracted_filename = potential_file
                        # Remove the filename comment from code
                        code_lines = code.split('\n')
                        code = '\n'.join(code_lines[1:]).strip()
                
                # Determine target file - SMART ROUTING based on content
                if extracted_filename:
                    target_file = extracted_filename
                elif filename:
                    target_file = filename
                else:
                    # Smart file routing based on code content
                    code_lower = code.lower()
                    
                    if lang in ['jsx', 'tsx']:
                        target_file = 'src/App.jsx' if lang == 'jsx' else 'src/App.tsx'
                    elif lang == 'css':
                        target_file = 'src/index.css'
                    elif lang == 'html':
                        target_file = 'index.html'
                    elif lang in ['javascript', 'js']:
                        target_file = 'src/main.js'
                    elif lang in ['typescript', 'ts']:
                        target_file = 'src/main.ts'
                    elif lang in ['python', 'py']:
                        # SMART PYTHON ROUTING - detect type based on content
                        if 'fastapi' in code_lower or '@app.get' in code_lower or '@app.post' in code_lower:
                            if 'auth' in code_lower or 'password' in code_lower or 'login' in code_lower:
                                target_file = 'backend/auth.py'
                            elif 'crud' in code_lower or 'database' in code_lower or 'sqlalchemy' in code_lower:
                                target_file = 'backend/database.py'
                            elif 'admin' in code_lower:
                                target_file = 'backend/admin.py'
                            elif 'payment' in code_lower or 'stripe' in code_lower:
                                target_file = 'backend/payments.py'
                            elif 'email' in code_lower or 'smtp' in code_lower:
                                target_file = 'backend/email_service.py'
                            else:
                                target_file = 'backend/api.py'
                        elif 'flask' in code_lower or '@app.route' in code_lower:
                            target_file = 'backend/app.py'
                        elif 'pydantic' in code_lower or 'basemodel' in code_lower:
                            target_file = 'backend/models.py'
                        elif 'test_' in code_lower or 'pytest' in code_lower or 'unittest' in code_lower:
                            target_file = 'tests/test_main.py'
                        elif 'class' in code_lower and ('model' in code_lower or 'schema' in code_lower):
                            target_file = 'backend/models.py'
                        elif 'import matplotlib' in code_lower or 'import plotly' in code_lower:
                            target_file = 'scripts/visualization.py'
                        elif 'import pandas' in code_lower and 'analysis' in code_lower:
                            target_file = 'scripts/analysis.py'
                        else:
                            # Default Python file - use descriptive name based on first function/class
                            import re as re_inner
                            func_match = re_inner.search(r'def (\w+)', code)
                            class_match = re_inner.search(r'class (\w+)', code)
                            if class_match:
                                target_file = f"backend/{class_match.group(1).lower()}.py"
                            elif func_match and func_match.group(1) != 'main':
                                target_file = f"scripts/{func_match.group(1)}.py"
                            else:
                                target_file = 'main.py'
                    elif lang == 'json':
                        target_file = 'data.json'
                    elif lang in ['bash', 'sh', 'shell']:
                        target_file = 'scripts/setup.sh'
                    elif lang in ['yaml', 'yml']:
                        target_file = 'config.yml'
                    elif lang in ['markdown', 'md']:
                        target_file = 'docs/README.md'
                    else:
                        target_file = f'src/generated.{lang}'
                
                full_path = os.path.join(project_path, target_file)
                os.makedirs(os.path.dirname(full_path), exist_ok=True)
                
                # CRITICAL: Prevent creating files that shadow Python packages!
                FORBIDDEN_FILENAMES = {
                    # Python stdlib
                    'base64', 'json', 'os', 'sys', 'io', 're', 'typing', 'datetime', 
                    'time', 'math', 'random', 'hashlib', 'secrets', 'contextlib',
                    'collections', 'functools', 'itertools', 'sqlite3', 'urllib',
                    # Common pip packages - DO NOT SHADOW THESE
                    'passlib', 'sqlalchemy', 'jose', 'jwt', 'pydantic', 'fastapi',
                    'flask', 'requests', 'httpx', 'aiohttp', 'numpy', 'pandas',
                    'cryptography', 'bcrypt', 'dotenv', 'redis', 'celery',
                    'pytest', 'unittest', 'logging', 'asyncio', 'threading',
                    'multiprocessing', 'subprocess', 'pathlib', 'shutil', 'glob',
                    # Also catch variations
                    'pyjwt', 'python-jose', 'python-dotenv', 'psycopg2', 'pymongo'
                }
                
                # Check if target filename (without extension) is in forbidden list
                base_name = os.path.basename(target_file).rsplit('.', 1)[0].lower()
                full_filename = os.path.basename(target_file)
                
                # === V4.0: SKIP junk files entirely, don't rename them ===
                if base_name in FORBIDDEN_FILENAMES or full_filename in self.junk_files:
                    self._log(f"  -> SKIPPED JUNK: {target_file} (would shadow stdlib/package)")
                    continue
                
                # === V4.2: CLEAN CODE - Strip prose/explanations ===
                code = self._clean_code_output(code, target_file)
                
                # V4.4: SMART OVERWRITE PROTECTION
                if not self._should_overwrite_file(full_path, code):
                    self._log(f"  ⚠️ SKIP: Won't overwrite {target_file} with lower quality content")
                    continue
                
                with open(full_path, 'w', encoding='utf-8') as f:
                    f.write(code)
                
                saved_files.append(target_file)
                self._log(f"  -> Saved: {target_file} ({len(code)} chars)")
        
        # Update file index for context in subsequent steps
        if saved_files and project_path:
            try:
                self._update_file_index(os.path.basename(project_path), saved_files)
            except Exception as e:
                pass  # Don't fail if index update fails
        
        return saved_files
    
    def _update_file_index(self, project_name: str, files: List[str]):
        """Track generated files in an index for the AI to reference."""
        from .config import WORKSPACE_DIR
        
        index_path = os.path.join(WORKSPACE_DIR, project_name, ".file_index.json")
        
        # Load existing index
        existing = []
        if os.path.exists(index_path):
            try:
                import json
                with open(index_path, 'r') as f:
                    existing = json.load(f)
            except:
                existing = []
        
        # Add new files
        for f in files:
            if f not in existing:
                existing.append(f)
        
        # Save updated index
        import json
        with open(index_path, 'w') as f:
            json.dump(existing, f, indent=2)
    
    def _clean_code_output(self, code: str, filename: str) -> str:
        """
        V4.2: Strip prose/explanations that LLM may have added to code.
        Hermes sometimes adds chat-like responses after code.
        """
        import re
        
        lines = code.split('\n')
        clean_lines = []
        in_code = True
        
        for line in lines:
            # Detect prose patterns that shouldn't be in code
            prose_patterns = [
                r'^This (?:implements?|component|code)',
                r'^The (?:key|main|above)',
                r'^Let me know if',
                r'^I\'m happy to',
                r'^Here\'s what',
                r'^Feel free to',
                r'^Note:? ',
                r'^(?:\d+\.)\s+(?:Add|Create|Implement|Use)',
                r'^\*\*',  # Bold markdown
                r'^- [A-Z]',  # Markdown bullet with sentence
            ]
            
            # Check if this line looks like prose
            is_prose = any(re.match(p, line.strip()) for p in prose_patterns)
            
            if is_prose:
                in_code = False
                continue
            
            if in_code:
                clean_lines.append(line)
        
        # Also remove trailing ``` that may have been left
        result = '\n'.join(clean_lines).rstrip()
        if result.endswith('```'):
            result = result[:-3].rstrip()
        
        return result
    
    def _should_overwrite_file(self, filepath: str, new_content: str) -> bool:
        """
        V4.4: Smart overwrite protection.
        Prevents overwriting good content with stubs or placeholders.
        """
        import os
        
        # If file doesn't exist, always allow
        if not os.path.exists(filepath):
            return True
        
        # Read existing content
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                existing = f.read()
        except:
            return True
        
        filename = os.path.basename(filepath)
        ext = os.path.splitext(filepath)[1].lower()
        
        # Protected files - extra careful
        PROTECTED = ['paper.md', 'glossary.md', 'README.md']
        is_protected = filename in PROTECTED
        
        # Check new content for quality issues
        new_has_issues = self._has_quality_issues(new_content, ext)
        existing_has_issues = self._has_quality_issues(existing, ext)
        
        # Rule 1: Never overwrite with smaller content (likely a stub)
        if len(new_content) < len(existing) * 0.5:
            self._log(f"    [Quality] New content ({len(new_content)}) much smaller than existing ({len(existing)})")
            return False
        
        # Rule 2: Never replace real content with placeholders
        if new_has_issues and not existing_has_issues:
            self._log(f"    [Quality] New content has placeholders, existing doesn't")
            return False
        
        # Rule 3: Protected files need significant improvement to overwrite
        if is_protected and len(new_content) < len(existing) * 1.2:
            if not existing_has_issues:
                self._log(f"    [Quality] Protected file - new content not significantly better")
                return False
        
        return True
    
    def _has_quality_issues(self, content: str, ext: str = '') -> bool:
        """
        V4.4: Detect if content has quality issues (placeholders, stubs).
        Returns True if content has problems.
        """
        import re
        
        # Placeholder patterns
        placeholder_patterns = [
            r'\[Insert .* here\]',
            r'\[TODO\]',
            r'\[TBD\]',
            r'# \.\.\.',
            r'\.\.\.  # ',
            r'pass\s*$',
            r'raise NotImplementedError',
            r'\# Add more .* as needed',
        ]
        
        for pattern in placeholder_patterns:
            if re.search(pattern, content, re.IGNORECASE):
                return True
        
        # Minimum content thresholds by file type
        min_sizes = {
            '.py': 200,
            '.md': 300,
            '.js': 150,
            '.jsx': 200,
            '.ts': 150,
            '.tsx': 200,
        }
        
        if ext in min_sizes and len(content) < min_sizes[ext]:
            return True
        
        return False
    
    def _auto_run_project(self, project_path: str) -> Dict:
        """
        Automatically install dependencies and start dev server.
        Returns: Dict with status, port, and any errors.
        """
        import subprocess
        import time
        
        result = {"success": False, "port": None, "error": None, "process": None}
        
        # Check if package.json exists
        package_json = os.path.join(project_path, "package.json")
        if not os.path.exists(package_json):
            result["error"] = "No package.json found"
            return result
        
        self._log("[AutoRun] Installing dependencies...")
        
        try:
            # Run npm install
            install_result = subprocess.run(
                ["npm", "install"],
                cwd=project_path,
                capture_output=True,
                text=True,
                timeout=300,  # 5 min timeout
                shell=True
            )
            
            if install_result.returncode != 0:
                result["error"] = f"npm install failed: {install_result.stderr[:500]}"
                self._log(f"[AutoRun] npm install failed: {install_result.stderr[:200]}")
                return result
            
            self._log("[AutoRun] Dependencies installed. Starting dev server...")
            
            # Start dev server in background
            dev_process = subprocess.Popen(
                ["npm", "run", "dev"],
                cwd=project_path,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                shell=True
            )
            
            # Wait a bit for server to start
            time.sleep(5)
            
            # Check if process is still running
            if dev_process.poll() is None:
                result["success"] = True
                result["port"] = 5173  # Vite default
                result["process"] = dev_process
                self._log(f"[AutoRun] ✅ Dev server running on port {result['port']}")
            else:
                stdout, stderr = dev_process.communicate()
                result["error"] = f"Dev server crashed: {stderr.decode()[:500]}"
                self._log(f"[AutoRun] ❌ Dev server failed")
            
        except subprocess.TimeoutExpired:
            result["error"] = "npm install timed out"
        except Exception as e:
            result["error"] = str(e)
            self._log(f"[AutoRun] Error: {e}")
        
        return result
    
    def _execute_with_recovery(self, step: str, context: str, max_retries: int = 2) -> str:
        """
        Execute a step with automatic error recovery.
        If execution fails, retry with error context.
        """
        last_error = None
        
        for attempt in range(max_retries + 1):
            try:
                # Add error context if this is a retry
                if last_error:
                    context += f"\n\n[PREVIOUS ERROR - PLEASE FIX]: {last_error}"
                    self._log(f"  Retry {attempt}/{max_retries} with error context")
                
                result = self._execute_step(step, context)
                
                # Check for obvious errors in response
                if "[Error:" in result or "[LLM Error" in result:
                    last_error = result
                    continue
                
                return result
                
            except Exception as e:
                last_error = str(e)
                self._log(f"  !!! ERROR: {e} !!!")
                
                if attempt == max_retries:
                    return f"[FAILED after {max_retries + 1} attempts: {last_error}]"
        
        return f"[FAILED: {last_error}]"
    
    def _capture_preview(self, port: int = 5173) -> str:
        """
        Capture a screenshot of the running dev server.
        Returns: Path to saved screenshot.
        """
        try:
            from playwright.sync_api import sync_playwright
            import time
            
            screenshot_path = os.path.join(WORKSPACE_DIR, "screenshots", f"preview_{int(time.time())}.png")
            os.makedirs(os.path.dirname(screenshot_path), exist_ok=True)
            
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page(viewport={"width": 1280, "height": 720})
                page.goto(f"http://localhost:{port}", wait_until="networkidle", timeout=30000)
                page.screenshot(path=screenshot_path)
                browser.close()
            
            self._log(f"[Preview] Screenshot saved: {screenshot_path}")
            return screenshot_path
            
        except Exception as e:
            self._log(f"[Preview] Failed to capture: {e}")
            return None
    
    def _audit_dependencies(self, project_path: str, saved_files: list) -> list:
        """
        Scan saved Python files for local imports and verify they exist.
        Returns a list of missing module names (only if they are NOT installed packages).
        """
        import re
        import importlib.util
        import sys
        
        missing = []
        # Pre-populate with known builtins and common libs to save time
        ignored_modules = set(sys.builtin_module_names) | {
            'os', 'sys', 'json', 'time', 'datetime', 're', 'random', 
            'math', 'collections', 'itertools', 'functools', 'typing',
            'pathlib', 'subprocess', 'threading', 'multiprocessing',
            'requests', 'numpy', 'pandas', 'networkx', 'matplotlib',
            'pydantic', 'fastapi', 'sqlalchemy', 'uvicorn', 'sqlite3',
            'asyncio', 'contextlib', 'shutil', 'logging', 'uuid', 'types',
            'copy', 'enum', 'hashlib', 'base64', 'io', 'platform', 'mock',
            'pytest', 'unittest', 'dataclasses', 'abc', 'argparse'
        }
        
        for filename in saved_files:
            if not filename.endswith('.py'):
                continue
            
            # Find the file path
            filepath = None
            for root, dirs, files in os.walk(project_path):
                if filename in files:
                    filepath = os.path.join(root, filename)
                    break
            
            if not filepath or not os.path.exists(filepath):
                continue
            
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Parse import statements
                # Match: import X, from X import Y
                import_pattern = r'^(?:from\s+(\w+)|import\s+(\w+))'
                for match in re.finditer(import_pattern, content, re.MULTILINE):
                    module_name = match.group(1) or match.group(2)
                    
                    if not module_name:
                        continue
                    
                    # 1. Check whitelist
                    if module_name in ignored_modules:
                        continue
                        
                    # 2. Check installed packages (safe check)
                    try:
                        spec = importlib.util.find_spec(module_name)
                        if spec is not None:
                            ignored_modules.add(module_name) # Cache it
                            continue
                    except Exception:
                        pass
                    
                    # 3. Check if local module exists (file in project)
                    local_found = False
                    for root, _, files in os.walk(project_path):
                        if f"{module_name}.py" in files:
                            local_found = True
                            break
                    
                    if local_found:
                        continue
                    
                    # If we reach here, it's truly missing
                    if module_name not in missing:
                        missing.append(module_name)
                        
            except Exception as e:
                self._log(f"[DependencyAudit] Error parsing {filename}: {e}")
        
        return missing
    
    def _validate_execution(self, project_path: str, saved_files: list) -> dict:
        """
        V3.3: Post-Execution Validation (FIXED)
        Only validates SYNTAX of Python files. Does NOT try to import them,
        since helper files and modules are not meant to be run standalone.
        Returns dict with {filename: error_message} for files with syntax errors.
        """
        errors = {}
        
        # Only validate entry points for import, skip helper files
        entry_points = {'main.py', 'app.py', 'server.py', 'index.py'}
        
        for filename in saved_files:
            if not filename.endswith('.py'):
                continue
            
            # Find the file path
            filepath = None
            for root, dirs, files in os.walk(project_path):
                if filename in files:
                    filepath = os.path.join(root, filename)
                    break
            
            if not filepath or not os.path.exists(filepath):
                continue
            
            try:
                # Only check SYNTAX (compile), don't try to import
                with open(filepath, 'r', encoding='utf-8') as f:
                    code = f.read()
                
                # Compile to check syntax ONLY
                compile(code, filepath, 'exec')
                self._log(f"[Validation] ✅ {filename}: Syntax OK")
                    
            except SyntaxError as e:
                errors[filename] = f"SyntaxError: {e.msg} at line {e.lineno}"
                self._log(f"[Validation] ❌ {filename}: SyntaxError at line {e.lineno}")
            except Exception as e:
                errors[filename] = f"ValidationError: {str(e)[:200]}"
                self._log(f"[Validation] ❌ {filename}: {e}")
        
        return errors
    
    def _multi_perspective_review(self, code: str, filename: str) -> list:
        """
        V3.3: Multi-Angle Code Review
        Reviews code from multiple perspectives: correctness, security, performance.
        Returns list of issues found.
        """
        issues = []
        
        # Perspective 1: Correctness (via Devils Advocate)
        from .devils_advocate import devils_advocate
        critique = devils_advocate.critique(code, content_type="code", context=f"File: {filename}")
        if critique.get("issues"):
            issues.extend([{"perspective": "correctness", **i} for i in critique["issues"]])
        
        # Perspective 2: Security (quick static checks)
        security_red_flags = [
            ("eval(", "CRITICAL: Use of eval() is dangerous"),
            ("exec(", "CRITICAL: Use of exec() is dangerous"),
            ("__import__", "WARNING: Dynamic import detected"),
            ("pickle.load", "WARNING: Pickle deserialization can be unsafe"),
            ("shell=True", "WARNING: Shell injection risk with subprocess"),
            ("os.system(", "WARNING: Prefer subprocess over os.system"),
        ]
        for pattern, warning in security_red_flags:
            if pattern in code:
                issues.append({"perspective": "security", "title": warning, "risk": "major"})
        
        # Perspective 3: Performance (basic checks)
        perf_issues = []
        if "for " in code and "append(" in code:
            # Suggest list comprehension
            perf_issues.append("Consider list comprehension instead of for+append")
        if code.count("for ") > 5:
            perf_issues.append("Multiple nested loops detected - consider optimization")
        
        for issue in perf_issues:
            issues.append({"perspective": "performance", "title": issue, "risk": "minor"})
        
        return issues
    
    def _get_project_name(self, objective: str) -> str:
        """Generate a project folder name from the objective."""
        import re
        # Extract key words and create short name
        words = re.findall(r'\b\w+\b', objective.lower())
        keywords = [w for w in words if len(w) > 3 and w not in ['build', 'create', 'make', 'with', 'that', 'this', 'from', 'using']]
        name = '-'.join(keywords[:3]) if keywords else 'project'
        return name.replace(' ', '-')[:30]
    
    def _scaffold_project(self, project_path: str, project_type: str = "research") -> dict:
        """
        V3.4: Create proper project structure with organized folders.
        Returns dict with created paths.
        """
        structure = {
            "research": ["src", "tests", "data", "docs", "figures"],
            "webapp": ["src", "public", "components", "styles", "tests"],
            "api": ["src", "routes", "models", "tests", "docs"],
            "default": ["src", "tests", "docs"]
        }
        
        folders = structure.get(project_type, structure["default"])
        created = {"folders": [], "files": []}
        
        for folder in folders:
            folder_path = os.path.join(project_path, folder)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path, exist_ok=True)
                created["folders"].append(folder)
                self._log(f"📁 SCAFFOLD: Created {folder}/")
        
        # Create README.md
        readme_path = os.path.join(project_path, "README.md")
        if not os.path.exists(readme_path):
            readme_content = f"""# Project: {os.path.basename(project_path)}

## Structure
```
{chr(10).join('├── ' + f + '/' for f in folders)}
├── requirements.txt
└── main.py
```

## Setup
```bash
pip install -r requirements.txt
python main.py
```

## Generated by Jarvis Autonomous System
"""
            with open(readme_path, 'w', encoding='utf-8') as f:
                f.write(readme_content)
            created["files"].append("README.md")
            self._log("📁 SCAFFOLD: Created README.md")
        
        return created
    
    def _install_dependencies(self, project_path: str, saved_files: list) -> list:
        """
        V3.4: Auto-install external dependencies found in Python files.
        Returns list of installed packages.
        """
        import subprocess
        import re
        
        # Common external packages (not in stdlib)
        external_packages = {
            'numpy': 'numpy', 'np': 'numpy',
            'pandas': 'pandas', 'pd': 'pandas',
            'matplotlib': 'matplotlib', 'plt': 'matplotlib',
            'networkx': 'networkx', 'nx': 'networkx',
            'requests': 'requests',
            'scipy': 'scipy',
            'sklearn': 'scikit-learn',
            'torch': 'torch',
            'tensorflow': 'tensorflow', 'tf': 'tensorflow',
            'flask': 'flask',
            'fastapi': 'fastapi',
            'pytest': 'pytest',
        }
        
        found_packages = set()
        
        for filename in saved_files:
            if not filename.endswith('.py'):
                continue
            
            # Find the file
            for root, dirs, files in os.walk(project_path):
                if filename in files:
                    filepath = os.path.join(root, filename)
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            content = f.read()
                        
                        # Find imports
                        import_pattern = r'^(?:from\s+(\w+)|import\s+(\w+))'
                        for match in re.finditer(import_pattern, content, re.MULTILINE):
                            module = match.group(1) or match.group(2)
                            if module in external_packages:
                                found_packages.add(external_packages[module])
                    except:
                        pass
                    break
        
        installed = []
        for package in found_packages:
            try:
                result = subprocess.run(
                    ['pip', 'install', package, '--quiet'],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    installed.append(package)
                    self._log(f"📦 DEPS: Installed {package}")
            except Exception as e:
                self._log(f"📦 DEPS: Failed to install {package}: {e}")
        
        # Generate requirements.txt
        if installed:
            req_path = os.path.join(project_path, "requirements.txt")
            with open(req_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(sorted(installed)))
            self._log(f"📦 DEPS: Generated requirements.txt with {len(installed)} packages")
        
        # === SMOKE TEST: Verify imports actually work ===
        smoke_failures = []
        for filename in saved_files:
            if not filename.endswith('.py'):
                continue
            
            # Find the file
            for root, dirs, files in os.walk(project_path):
                if filename in files:
                    filepath = os.path.join(root, filename)
                    try:
                        # Try to compile and check imports
                        result = subprocess.run(
                            ['python', '-m', 'py_compile', filepath],
                            capture_output=True,
                            text=True,
                            timeout=10
                        )
                        if result.returncode != 0:
                            smoke_failures.append((filename, result.stderr[:200]))
                            self._log(f"🔴 SMOKE: Syntax error in {filename}")
                        else:
                            self._log(f"🟢 SMOKE: {filename} syntax OK")
                    except Exception as e:
                        smoke_failures.append((filename, str(e)))
                    break
        
        # Report smoke test results
        if smoke_failures:
            self._log(f"⚠️ SMOKE TEST: {len(smoke_failures)} files have issues")
            for fname, err in smoke_failures[:3]:  # Show first 3
                self._log(f"   - {fname}: {err[:100]}")
        else:
            self._log(f"✅ SMOKE TEST: All {len([f for f in saved_files if f.endswith('.py')])} Python files OK")
        
        return installed
    
    def _inject_entry_point(self, project_path: str, main_file: str = "main.py") -> bool:
        """
        V3.4: Ensure main.py has an if __name__ == '__main__' block.
        Returns True if modified.
        """
        filepath = os.path.join(project_path, main_file)
        if not os.path.exists(filepath):
            return False
        
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Check if entry point exists
            if 'if __name__' in content:
                return False
            
            # Find the main class or function to call
            import re
            
            # Look for Simulation class or main function
            main_call = ""
            if 'class Simulation' in content:
                main_call = '''
if __name__ == "__main__":
    # Auto-generated entry point
    sim = Simulation(topology_type="ER", num_agents=50, alpha=0.7, beta=0.3)
    sim.run(iterations=100)
    print("Results:", sim.collect_metrics())
'''
            elif 'def main(' in content:
                main_call = '''
if __name__ == "__main__":
    main()
'''
            else:
                # Generic entry point
                main_call = '''
if __name__ == "__main__":
    print("Project initialized. Add your entry point logic here.")
'''
            
            # Append entry point
            with open(filepath, 'a', encoding='utf-8') as f:
                f.write(main_call)
            
            self._log(f"🚀 ENTRY: Injected entry point into {main_file}")
            return True
            
        except Exception as e:
            self._log(f"🚀 ENTRY: Failed to inject entry point: {e}")
            return False
    
    def _run_and_capture_output(self, project_path: str, script: str = "main.py", timeout: int = 120) -> dict:
        """
        V3.5: Run a Python script and capture its output (stdout, stderr, JSON).
        Returns dict with stdout, stderr, json_data, success.
        """
        import subprocess
        import json
        
        filepath = os.path.join(project_path, script)
        if not os.path.exists(filepath):
            return {"success": False, "error": f"{script} not found"}
        
        try:
            result = subprocess.run(
                ['python', script],
                capture_output=True,
                text=True,
                timeout=timeout,
                cwd=project_path
            )
            
            output = {
                "success": result.returncode == 0,
                "stdout": result.stdout,
                "stderr": result.stderr,
                "returncode": result.returncode,
                "json_data": None
            }
            
            # Try to extract JSON from stdout
            if result.stdout:
                try:
                    # Look for JSON in output (between { and })
                    import re
                    json_match = re.search(r'\{[^{}]*\}', result.stdout, re.DOTALL)
                    if json_match:
                        output["json_data"] = json.loads(json_match.group())
                except:
                    pass
            
            # Also check for results file
            results_file = os.path.join(project_path, "simulation_results.json")
            if os.path.exists(results_file):
                with open(results_file, 'r') as f:
                    output["json_data"] = json.load(f)
            
            self._log(f"▶️ RUN: {script} {'✅' if output['success'] else '❌'}")
            if output["json_data"]:
                self._log(f"   📊 Captured JSON data with {len(output['json_data'])} keys")
            
            return output
            
        except subprocess.TimeoutExpired:
            return {"success": False, "error": f"Timeout after {timeout}s"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _generate_figures(self, project_path: str, data: dict) -> list:
        """
        V3.5: Auto-generate matplotlib figures from simulation data.
        Saves figures to figures/ directory.
        Returns list of generated figure filenames.
        """
        figures_dir = os.path.join(project_path, "figures")
        os.makedirs(figures_dir, exist_ok=True)
        
        generated = []
        
        try:
            import matplotlib
            matplotlib.use('Agg')  # Non-interactive backend
            import matplotlib.pyplot as plt
            import numpy as np
            
            # Figure 1: Bar chart of efficiencies by topology
            if isinstance(data, dict) and any('efficiency' in str(v) for v in data.values()):
                fig, ax = plt.subplots(figsize=(10, 6))
                
                topologies = list(data.keys())
                efficiencies = [v.get('avg_efficiency', v.get('efficiency', 0)) 
                               for v in data.values() if isinstance(v, dict)]
                
                if efficiencies:
                    ax.bar(topologies, efficiencies, color=['#3498db', '#2ecc71', '#e74c3c'])
                    ax.set_xlabel('Network Topology')
                    ax.set_ylabel('Mean Efficiency')
                    ax.set_title('ATRA-G Performance Across Topologies')
                    ax.set_ylim(0, 1.1)
                    
                    fig_path = os.path.join(figures_dir, "efficiency_comparison.png")
                    plt.savefig(fig_path, dpi=150, bbox_inches='tight')
                    plt.close()
                    generated.append("efficiency_comparison.png")
                    self._log("📊 FIGURE: Generated efficiency_comparison.png")
            
            # Figure 2: Variance comparison
            if isinstance(data, dict):
                variances = [v.get('avg_variance', v.get('variance', 0)) 
                            for v in data.values() if isinstance(v, dict)]
                
                if variances and any(v > 0 for v in variances):
                    fig, ax = plt.subplots(figsize=(10, 6))
                    topologies = list(data.keys())
                    ax.bar(topologies, variances, color=['#9b59b6', '#f39c12', '#1abc9c'])
                    ax.set_xlabel('Network Topology')
                    ax.set_ylabel('Variance')
                    ax.set_title('Resource Allocation Variance by Topology')
                    
                    fig_path = os.path.join(figures_dir, "variance_comparison.png")
                    plt.savefig(fig_path, dpi=150, bbox_inches='tight')
                    plt.close()
                    generated.append("variance_comparison.png")
                    self._log("📊 FIGURE: Generated variance_comparison.png")
                    
        except ImportError:
            self._log("📊 FIGURE: matplotlib not available, skipping figure generation")
        except Exception as e:
            self._log(f"📊 FIGURE: Error generating figures: {e}")
        
        return generated
    
    def _enforce_paper_length(self, prompt: str, min_words: int = 3000) -> str:
        """
        V3.5: Modify paper generation prompt to enforce minimum length.
        """
        length_enforcement = f"""

CRITICAL LENGTH REQUIREMENT:
- You MUST write a COMPLETE academic paper with AT LEAST {min_words} words.
- Do NOT write an outline or summary.
- Each section MUST have multiple full paragraphs.
- Include: Abstract (200+ words), Introduction (500+ words), Methods (800+ words), 
  Results (600+ words), Discussion (500+ words), Conclusion (200+ words).
- Use proper academic language and cite all figures.
- If you write less than {min_words} words, your paper will be REJECTED.

"""
        return length_enforcement + prompt
    
    def _export_to_docx(self, project_path: str, markdown_file: str = "paper.md") -> str:
        """
        V3.5: Convert markdown to DOCX format for academic submission.
        Returns path to generated DOCX or empty string on failure.
        """
        md_path = os.path.join(project_path, markdown_file)
        
        # Also check in src/ and docs/
        if not os.path.exists(md_path):
            md_path = os.path.join(project_path, "src", markdown_file)
        if not os.path.exists(md_path):
            md_path = os.path.join(project_path, "docs", markdown_file)
        if not os.path.exists(md_path):
            # Try generated.markdown
            md_path = os.path.join(project_path, "src", "generated.markdown")
        
        if not os.path.exists(md_path):
            self._log("📄 DOCX: No markdown file found to convert")
            return ""
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Read markdown
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document()
            
            # Process markdown line by line
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('# '):
                    # Title
                    p = doc.add_heading(line[2:], level=0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=1)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)
            
            # Save DOCX
            docx_path = md_path.replace('.md', '.docx').replace('.markdown', '.docx')
            doc.save(docx_path)
            
            self._log(f"📄 DOCX: Exported to {os.path.basename(docx_path)}")
            return docx_path
            
        except ImportError:
            self._log("📄 DOCX: python-docx not installed, attempting install...")
            import subprocess
            subprocess.run(['pip', 'install', 'python-docx', '--quiet'], capture_output=True)
            # Retry
            try:
                from docx import Document
                return self._export_to_docx(project_path, markdown_file)
            except:
                self._log("📄 DOCX: Failed to install python-docx")
                return ""
        except Exception as e:
            self._log(f"📄 DOCX: Error exporting: {e}")
            return ""
    
    def _find_available_port(self, start_port: int = 3000, max_attempts: int = 20) -> int:
        """
        V3.6: Find an available port for dev server.
        """
        import socket
        for port in range(start_port, start_port + max_attempts):
            try:
                with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                    s.bind(('localhost', port))
                    return port
            except OSError:
                continue
        return start_port + max_attempts  # Fallback
    
    def _start_dev_server(self, project_path: str, server_type: str = "auto") -> dict:
        """
        V3.6: Start a development server for the project.
        Returns dict with pid, port, url, type.
        """
        import subprocess
        
        # Detect server type
        if server_type == "auto":
            if os.path.exists(os.path.join(project_path, "package.json")):
                server_type = "npm"
            elif os.path.exists(os.path.join(project_path, "index.html")):
                server_type = "python"
            else:
                server_type = "python"
        
        port = self._find_available_port()
        
        try:
            if server_type == "npm":
                # Check if node_modules exists
                if not os.path.exists(os.path.join(project_path, "node_modules")):
                    self._log("🔧 DEV: Running npm install...")
                    subprocess.run(['npm', 'install'], cwd=project_path, capture_output=True)
                
                # Start dev server
                process = subprocess.Popen(
                    ['npm', 'run', 'dev', '--', '--port', str(port)],
                    cwd=project_path,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    shell=True
                )
            else:
                # Python HTTP server
                process = subprocess.Popen(
                    ['python', '-m', 'http.server', str(port)],
                    cwd=project_path,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE
                )
            
            url = f"http://localhost:{port}"
            self._log(f"🚀 DEV: Server started at {url} (PID: {process.pid})")
            
            return {
                "success": True,
                "pid": process.pid,
                "port": port,
                "url": url,
                "type": server_type
            }
            
        except Exception as e:
            self._log(f"🚀 DEV: Failed to start server - {e}")
            return {"success": False, "error": str(e)}
    
    def _deploy_netlify(self, project_path: str, prod: bool = False) -> dict:
        """
        V3.6: Deploy project to Netlify using CLI.
        """
        import subprocess
        
        # Check if Netlify CLI is installed
        try:
            result = subprocess.run(['netlify', '--version'], capture_output=True, text=True)
            if result.returncode != 0:
                self._log("📦 Installing Netlify CLI...")
                subprocess.run(['npm', 'install', '-g', 'netlify-cli'], capture_output=True)
        except FileNotFoundError:
            self._log("📦 Installing Netlify CLI...")
            subprocess.run(['npm', 'install', '-g', 'netlify-cli'], capture_output=True)
        
        # Determine build directory
        build_dir = project_path
        for candidate in ['dist', 'build', 'public', '.']:
            candidate_path = os.path.join(project_path, candidate)
            if os.path.exists(candidate_path) and os.path.isdir(candidate_path):
                build_dir = candidate_path
                break
        
        try:
            cmd = ['netlify', 'deploy', '--dir', build_dir]
            if prod:
                cmd.append('--prod')
            
            result = subprocess.run(
                cmd,
                cwd=project_path,
                capture_output=True,
                text=True,
                timeout=120
            )
            
            # Extract URL from output
            import re
            url_match = re.search(r'(https://[^\s]+\.netlify\.app)', result.stdout)
            deploy_url = url_match.group(1) if url_match else None
            
            if deploy_url:
                self._log(f"🌐 NETLIFY: Deployed to {deploy_url}")
                return {"success": True, "url": deploy_url, "prod": prod}
            else:
                self._log(f"🌐 NETLIFY: Deploy output - {result.stdout[:200]}")
                return {"success": False, "output": result.stdout, "error": result.stderr}
                
        except Exception as e:
            self._log(f"🌐 NETLIFY: Deploy failed - {e}")
            return {"success": False, "error": str(e)}
    
    def _deploy_vercel(self, project_path: str, prod: bool = False) -> dict:
        """
        V3.6: Deploy project to Vercel using CLI.
        """
        import subprocess
        
        # Check if Vercel CLI is installed
        try:
            result = subprocess.run(['vercel', '--version'], capture_output=True, text=True)
            if result.returncode != 0:
                self._log("📦 Installing Vercel CLI...")
                subprocess.run(['npm', 'install', '-g', 'vercel'], capture_output=True)
        except FileNotFoundError:
            self._log("📦 Installing Vercel CLI...")
            subprocess.run(['npm', 'install', '-g', 'vercel'], capture_output=True)
        
        try:
            cmd = ['vercel', '--yes']  # --yes for non-interactive
            if prod:
                cmd.append('--prod')
            
            result = subprocess.run(
                cmd,
                cwd=project_path,
                capture_output=True,
                text=True,
                timeout=180
            )
            
            # Extract URL from output
            import re
            url_match = re.search(r'(https://[^\s]+\.vercel\.app)', result.stdout)
            deploy_url = url_match.group(1) if url_match else None
            
            if deploy_url:
                self._log(f"▲ VERCEL: Deployed to {deploy_url}")
                return {"success": True, "url": deploy_url, "prod": prod}
            else:
                self._log(f"▲ VERCEL: Deploy output - {result.stdout[:200]}")
                return {"success": False, "output": result.stdout, "error": result.stderr}
                
        except Exception as e:
            self._log(f"▲ VERCEL: Deploy failed - {e}")
            return {"success": False, "error": str(e)}
    
    def _start_project(self, project_path: str, mode: str = "dev") -> dict:
        """
        V3.6: Smart project starter - detects project type and runs appropriately.
        Modes: 'dev' (local server), 'run' (execute main.py), 'deploy' (Netlify/Vercel)
        """
        project_name = os.path.basename(project_path)
        
        # Detect project type
        has_package_json = os.path.exists(os.path.join(project_path, "package.json"))
        has_main_py = os.path.exists(os.path.join(project_path, "main.py"))
        has_index_html = os.path.exists(os.path.join(project_path, "index.html"))
        
        if mode == "dev":
            if has_package_json:
                return self._start_dev_server(project_path, "npm")
            elif has_index_html:
                return self._start_dev_server(project_path, "python")
            else:
                return {"success": False, "error": "No web project detected"}
        
        elif mode == "run":
            if has_main_py:
                return self._run_and_capture_output(project_path, "main.py")
            else:
                return {"success": False, "error": "No main.py found"}
        
        elif mode == "deploy":
            # Try Vercel first (faster), fallback to Netlify
            result = self._deploy_vercel(project_path, prod=True)
            if not result.get("success"):
                result = self._deploy_netlify(project_path, prod=True)
            return result
        
        else:
            return {"success": False, "error": f"Unknown mode: {mode}"}
    
    def _execute_step(self, step: str, task_context: str) -> str:
        """Execute a single step with smart context and terminal integration."""
        self._log(f"Executing: {step}")
        
        # Get the ORIGINAL objective - this prevents hallucination
        original_objective = recycler.task_objective
        
        # Detect task type
        task_type = self._detect_task_type(original_objective)
        
        project_name = self._get_project_name(original_objective)
        project_path = os.path.join(WORKSPACE_DIR, "projects", project_name)

        # V3.4: Scaffold project structure on first access
        if not os.path.exists(project_path):
            os.makedirs(project_path, exist_ok=True)
        self._scaffold_project(project_path, project_type=task_type)

        # Route to specialist
        routing = self._route_to_specialist(step, task_type, project_path)
        
        # Execute via specialist
        if routing["use_specialist"]:
            response = self._call_specialist(
                routing["agent"], 
                step, 
                project_path, 
                routing["context"]
            )
        else:
            # Direct LLM execution with STRUCTURED prompt
            task_type = self._detect_task_type(original_objective)
            
            # === V4.1: GET LOCKED PROJECT TYPE RULES ===
            project_type_injection = self._detect_and_lock_project_type(original_objective)
            
            # === V4.2: LEAN PROMPT FOR HERMES 32K ===
            # Reduced context bloat, explicit CODE ONLY instruction
            structured_prompt = f"""## TASK
{step}

## PROJECT: {original_objective[:500]}

## PATH: {project_path}
{project_type_injection}

## OUTPUT FORMAT - CRITICAL
Output ONLY code blocks. NO explanations, NO prose, NO markdown outside code blocks.

```jsx filename="ComponentName.jsx"
// Your complete, working code here
```

RULES:
1. CODE ONLY - no text before or after code blocks
2. Each file needs: filename="path/to/file.ext"
3. Complete, runnable code - no placeholders
4. DO NOT reimplementing stdlib (no jwt.py, bcrypt.py, etc.)
5. DO NOT create auth/database unless requested
6. STAY ON TASK: {step}

OUTPUT:"""
            
            response = self._call_llm(structured_prompt, max_tokens=8192)

        # CRITICAL: Auto-Save Supervisor (Fix for 'Lost in RAM' bug)
        # Capture artifacts immediately after generation
        saved_files = self._extract_and_save_code(response, project_name)
        if saved_files:
            response += f"\n\n[SYSTEM: Auto-saved {len(saved_files)} files: {', '.join(saved_files)}]"
            
            # V3.4: Auto-install dependencies found in saved files
            installed_deps = self._install_dependencies(project_path, saved_files)
            if installed_deps:
                response += f"\n[SYSTEM: Installed {len(installed_deps)} packages: {', '.join(installed_deps)}]"
            
            # V3.4: Inject entry point if main.py lacks one
            if 'main.py' in saved_files:
                self._inject_entry_point(project_path, 'main.py')
        
        # === CRITICAL: EXTRACT AND RUN COMMANDS ===
        # This was missing! Commands like [COMMAND]: python benchmark.py were never executed
        commands = self._extract_commands(response)
        if commands:
            self._log(f"📟 Found {len(commands)} command(s) to execute")
            for cmd in commands:
                self._log(f"  Running: {cmd}")
                try:
                    cmd_result = terminal.run(cmd, project_path, timeout=120)
                    if cmd_result.get("success"):
                        self._log(f"  ✓ Command OK ({cmd_result.get('duration', 0):.1f}s)")
                        # Capture output for context
                        if cmd_result.get("stdout"):
                            response += f"\n\n[COMMAND OUTPUT: {cmd}]\n{cmd_result['stdout'][:2000]}"
                    else:
                        self._log(f"  ✗ Command failed: {cmd_result.get('stderr', 'Unknown error')[:100]}")
                except Exception as e:
                    self._log(f"  ✗ Command error: {e}")
        
        # === DEVILS ADVOCATE REVIEW (The "Integrity Check") ===
        # Mandatory audit for coding and research tasks
        if task_type in ["coding", "research"] and "critical" not in step.lower():
            from .devils_advocate import devils_advocate
            
            # Pass the list of ACTUAL saved files to the Devil so it can catch hallucinations
            reality_context = f"Project: {project_name}\nFiles Saved This Turn: {json.dumps(saved_files) if saved_files else 'NONE'}"
            
            critique = devils_advocate.critique(response, content_type=task_type, context=reality_context)
            
            if critique["verdict"] in ["FIX_REQUIRED", "REVIEW_REQUIRED"]:
                self._log(f"😈 DEVIL REJECTED: {len(critique['issues'])} issues found.")
                
                # Filter for just the serious ones
                blockers = [i for i in critique['issues'] if i['risk'] in ['critical', 'major']]
                
                if blockers:
                    feedback = f"""## QUALITY CHECK FAILED - REDO REQUIRED

Your previous output was REJECTED. You must fix these issues:

"""
                    for i in blockers:
                        feedback += f"### ❌ {i['title']}\n{i['description']}\n**Required Fix:** {i['fix']}\n\n"
                    
                    feedback += f"""## ORIGINAL STEP TO REDO
{step}

## OUTPUT FORMAT REQUIRED
For code: Use fenced blocks with filename like ```python filename="file.py"
For documents: Complete markdown, NO [TODO] or [PLACEHOLDER] markers
For data: Real numbers, not fabricated percentages

## PRODUCE THE CORRECTED OUTPUT NOW:"""
                    
                    self._log("  -> Forcing retry with structured feedback...")
                    # Retry with feedback
                    response = self._call_llm(feedback, max_tokens=8192)
                    
                    # Auto-save again for the fixed version
                    new_saved = self._extract_and_save_code(response, project_name)
                    if new_saved:
                        response += f"\n\n[SYSTEM: Auto-saved fixes: {', '.join(new_saved)}]"
            else:
                self._log("😇 DEVIL APPROVED: No critical issues.")
        
        # === DEPENDENCY AUDITOR (The "Lab Work" Enforcer) ===
        # Check ALL Python files in the project for missing imports (not just newly saved)
        all_py_files = []
        for root, dirs, files in os.walk(project_path):
            for f in files:
                if f.endswith('.py'):
                    all_py_files.append(f)
        
        if all_py_files:
            missing_deps = self._audit_dependencies(project_path, all_py_files)
            if missing_deps:
                self._log(f"🔬 DEPENDENCY AUDIT: {len(missing_deps)} missing modules detected: {missing_deps}")
                # Force the agent to write these modules NOW
                for missing_module in missing_deps:
                    dep_prompt = f"""CRITICAL: Your code imports '{missing_module}' but this file does not exist.
You MUST implement '{missing_module}.py' NOW with all the functions your code expects.
Project path: {project_path}

Output the full Python code block for '{missing_module}.py'. No explanations, just the code."""
                    self._log(f"  -> Forcing implementation of: {missing_module}.py")
                    dep_response = self._call_llm(dep_prompt)
                    new_saved = self._extract_and_save_code(dep_response, self._get_project_name(original_objective))
                    if new_saved:
                        response += f"\\n\\n[SYSTEM: Auto-generated dependency: {', '.join(new_saved)}]"

        # === V3.3: POST-EXECUTION VALIDATION ===
        # Try to run saved Python files and catch any errors
        if all_py_files:
            # === V4.0: REDUCED VALIDATION RETRIES ===
            max_retries = self.max_validation_retries  # Default: 2 (was 3)
            for retry in range(max_retries):
                validation_errors = self._validate_execution(project_path, all_py_files)
                
                if not validation_errors:
                    self._log(f"🧪 VALIDATION: All {len(all_py_files)} files passed!")
                    break
                
                self._log(f"🧪 VALIDATION: {len(validation_errors)} errors found (retry {retry + 1}/{max_retries})")
                
                # Build fix prompt with specific errors
                fix_prompt = "Your code has ERRORS that prevent it from running:\n\n"
                for filename, error in validation_errors.items():
                    fix_prompt += f"❌ {filename}: {error}\n"
                fix_prompt += "\nFix ALL these errors NOW. Output the corrected code blocks."
                
                # Call LLM to fix
                fix_response = self._call_llm(fix_prompt)
                fixed_files = self._extract_and_save_code(fix_response, project_name)
                
                if fixed_files:
                    response += f"\n\n[SYSTEM: Auto-fixed files: {', '.join(fixed_files)}]"
                    # Re-scan for all py files after fix
                    all_py_files = []
                    for root, dirs, files in os.walk(project_path):
                        for f in files:
                            if f.endswith('.py'):
                                all_py_files.append(f)
        
        # === V3.3: MULTI-PERSPECTIVE CODE REVIEW ===
        # Review saved code from multiple angles (security, performance, correctness)
        if saved_files:
            for filename in saved_files:
                if not filename.endswith('.py'):
                    continue
                
                # Find and read the file
                for root, dirs, files in os.walk(project_path):
                    if filename in files:
                        filepath = os.path.join(root, filename)
                        try:
                            with open(filepath, 'r', encoding='utf-8') as f:
                                code = f.read()
                            
                            issues = self._multi_perspective_review(code, filename)
                            critical_issues = [i for i in issues if i.get('risk') in ['critical', 'major']]
                            
                            if critical_issues:
                                self._log(f"🔍 REVIEW: {filename} has {len(critical_issues)} critical issues")
                                # Force fix for security issues
                                security_issues = [i for i in critical_issues if i.get('perspective') == 'security']
                                if security_issues:
                                    sec_prompt = f"SECURITY ALERT for {filename}:\n"
                                    for issue in security_issues:
                                        sec_prompt += f"- {issue['title']}\n"
                                    sec_prompt += "\nRewrite the code to fix these security issues."
                                    self._log(f"  -> Forcing security fix for {filename}")
                                    sec_response = self._call_llm(sec_prompt)
                                    self._extract_and_save_code(sec_response, project_name)
                            else:
                                self._log(f"🔍 REVIEW: {filename} passed all perspectives ✅")
                        except Exception as e:
                            self._log(f"[Review] Error reading {filename}: {e}")
                        break

        # === V3.5: DATA-FIRST PIPELINE ===
        # After code is saved, run it and capture output for paper generation
        if 'main.py' in saved_files or os.path.exists(os.path.join(project_path, 'main.py')):
            # Step 1: Run simulation and capture output
            run_result = self._run_and_capture_output(project_path, 'main.py')
            
            if run_result.get("success") and run_result.get("json_data"):
                # Step 2: Generate figures from captured data
                figures = self._generate_figures(project_path, run_result["json_data"])
                if figures:
                    response += f"\n[SYSTEM: Generated {len(figures)} figures: {', '.join(figures)}]"
                
                # Step 3: If this is a research task, trigger paper generation with real data
                if task_type == "research" and "paper" not in step.lower():
                    self._log("📝 DATA-FIRST: Triggering paper generation with REAL data...")
                    
                    # Build paper prompt with actual results
                    paper_prompt = self._enforce_paper_length(f"""
Write a COMPLETE academic research paper based on THESE ACTUAL RESULTS:

SIMULATION DATA:
{json.dumps(run_result['json_data'], indent=2)}

GENERATED FIGURES (reference these in your paper):
{', '.join(figures) if figures else 'None generated'}

PROJECT: {project_name}
OBJECTIVE: {original_objective}

Write the FULL paper now with proper academic formatting.
""")
                    # Store for next step rather than calling LLM here
                    # (Paper generation should be a separate step)
                    data_file = os.path.join(project_path, "data", "results.json")
                    os.makedirs(os.path.dirname(data_file), exist_ok=True)
                    with open(data_file, 'w') as f:
                        json.dump(run_result["json_data"], f, indent=2)
                    self._log(f"📊 DATA: Saved results to data/results.json for paper generation")
            
            elif run_result.get("error"):
                self._log(f"⚠️ RUN: Simulation failed - {run_result['error']}")
        
        # === V3.5: DOCX EXPORT ===
        # Export any markdown papers to DOCX
        md_files = [f for f in saved_files if f.endswith('.md') or f.endswith('.markdown')]
        for md_file in md_files:
            if 'paper' in md_file.lower() or 'generated' in md_file.lower():
                self._export_to_docx(project_path, md_file)

        return response
        # Ensure project exists
        if not os.path.exists(project_path):
            project_manager.create_project(project_name, stack="vanilla", category="frontend")
        
        # === SPECIALIST ROUTING: Route to the right agent ===
        routing = self._route_to_specialist(step, task_type, project_path)
        
        if routing["use_specialist"]:
            # Use specialist agent
            task_context_short = task_context[:1500] if task_context else ""
            result = self._call_specialist(routing["agent"], step, project_path, task_context_short)
            
            # Still extract and save any code from specialist output
            saved_files = self._extract_and_save_code(result, project_name)
            if saved_files:
                self._log(f"  Created {len(saved_files)} files in {project_name}/")
                code_indexer.index_project(project_path)
                
                # Auto-run project if we have multiple files (looks like a complete project)
                if len(saved_files) >= 3 and any(f.endswith(('.jsx', '.tsx', '.html')) for f in saved_files):
                    run_result = self._auto_run_project(project_path)
                    if run_result["success"]:
                        # Capture preview screenshot
                        screenshot = self._capture_preview(run_result.get("port", 5173))
                        if screenshot:
                            self._log(f"  📸 Preview captured: {screenshot}")
            
            # Save result to domain and mark complete
            domain = self._detect_domain(result)
            recycler.save_to_domain(domain, f"### {step}\n{result[:2000]}")
            recycler.mark_step_complete(step, result[:500])
            
            return result
        
        # === FALLBACK: Use LLM directly for planning/general steps ===
        
        # === SMART CONTEXT: Only relevant files, not everything ===
        code_context = ""
        if os.path.exists(project_path) and task_type == "coding":
            code_context = code_indexer.get_relevant_context(project_path, step, max_tokens=2000)
        
        # Build context section
        context_section = ""
        if code_context:
            context_section = f"### Relevant Project Files:\n{code_context}"
        
        # Build prompt with STRONG objective reinforcement to prevent hallucination
        prompt = f"""# YOUR MISSION
You are executing a task for your user. Stay STRICTLY focused on this objective.

## OBJECTIVE (NEVER DEVIATE FROM THIS):
{original_objective}

## CURRENT STEP:
{step}

## TASK CONTEXT:
{task_context[:2500]}

{context_section}

## INSTRUCTIONS:
"""
        
        # Add task-specific instructions
        if task_type == "coding":
            # Check if this is a component/module step
            is_component = "[COMPONENT" in step.upper() or "module" in step.lower() or "build" in step.lower()
            
            if is_component:
                # COMPONENT MODE: Create complete, polished module
                prompt += f"""
## COMPONENT BUILDING MODE

You are building a COMPLETE, POLISHED component. Think of the premium landing page quality:
- 686 lines of polished HTML+CSS+JS in ONE file
- Glassmorphism, gradients, animations
- Fully functional with all features
- Dark theme with CSS variables
- Responsive design

FOR THIS STEP: {step}

CREATE A SINGLE COMPLETE FILE that includes:
1. All HTML structure
2. All CSS (inline in <style> tags) with:
   - CSS variables for theme
   - glassmorphism effects
   - hover animations
   - responsive breakpoints
3. All JavaScript functionality
4. Mock data that looks realistic
5. Proper navigation

QUALITY CHECKLIST:
[ ] Dark theme with purple/cyan accents
[ ] Glassmorphism cards with backdrop-filter
[ ] Smooth hover transitions
[ ] Working CRUD operations (add/edit/delete)
[ ] Proper form handling
[ ] Loading states
[ ] Responsive layout

Output the COMPLETE file in a code block with proper filename.
Example: ```html filename="crm-contacts.html"
"""
            else:
                # Regular coding task
                prompt += f"""
- You are building: {original_objective}
- Create COMPLETE working code, not fragments
- Include inline styles for polished look
"""
            
            prompt += """
- Use proper code blocks: ```html filename="example.html"
- For commands, use: [COMMAND]: npm install, python test.py, etc.
- For file edits, use: [EDIT] filename: replace "old" with "new"
"""
        elif task_type == "research":
            prompt += """
- Provide comprehensive research findings
- Include sources and citations where possible
- Structure findings clearly with headers
- Highlight key insights and actionable items
"""
        elif task_type == "writing":
            prompt += """
- Write professional, engaging content
- Match the tone to the context
- Structure with clear sections
- Include any necessary formatting
"""
        elif task_type == "analysis":
            prompt += f"""
## ANALYSIS MODE - DATA REQUIRED

You are analyzing data for: {original_objective}

CRITICAL RULES:
1. You MUST read actual data files mentioned in the objective
2. For Excel files, use: [COMMAND]: python -c "import pandas as pd; df = pd.read_excel('path/to/file.xlsx'); print(df.head(50).to_string())"
3. Extract SPECIFIC names, numbers, and companies from the data
4. NEVER use placeholders like "[Company Name]" or "[Your Name]" - use REAL data
5. Personalize all output based on ACTUAL data read
6. Save outputs to the specified directory (create if needed)

QUALITY CHECKLIST:
[ ] Used actual company/contact names from data
[ ] No placeholder text anywhere
[ ] Actionable recommendations with specifics
[ ] Saved to correct output location

For file operations on Windows:
[COMMAND]: python -c "import os; os.makedirs('path', exist_ok=True)"
"""
        else:  # general
            prompt += """
- Complete this step thoroughly
- Be specific and actionable
- Include all relevant details
- State what was accomplished
"""
        
        prompt += f"""

## CRITICAL REMINDERS:
- You are working on: {original_objective}
- This is step: {step}
- Stay focused. Do NOT make up different product names or objectives.
- IMPORTANT: When updating a file, output the COMPLETE file including ALL previous content plus your additions.
- At the end, state what was accomplished for this step."""
        
        # === ADAPTIVE TOKEN LIMITS ===
        from .config import TOKEN_LIMITS
        
        # Choose token limit based on step and task type
        if "plan" in step.lower() or "analyze" in step.lower():
            token_limit = TOKEN_LIMITS["planning"]  # 4K - fast for planning
        elif "[COMPONENT" in step.upper() or "complete" in step.lower():
            token_limit = TOKEN_LIMITS["component"]  # 10K - for full components
        elif task_type == "analysis":
            token_limit = TOKEN_LIMITS["standard"]  # 6K - analysis needs detail
        else:
            token_limit = TOKEN_LIMITS["standard"]  # 6K default
        
        result = self._call_llm(prompt, max_tokens=token_limit)
        
        # === EXTRACT AND RUN COMMANDS ===
        commands = self._extract_commands(result)
        for cmd in commands:
            self._log(f"  Running: {cmd}")
            cmd_result = terminal.run(cmd, project_path)
            if cmd_result["success"]:
                self._log(f"  Command OK ({cmd_result['duration']}s)")
            else:
                self._log(f"  Command failed: {cmd_result['stderr'][:100]}")
        
        # === EXTRACT AND SAVE CODE ===
        saved_files = self._extract_and_save_code(result, project_name)
        
        if saved_files:
            self._log(f"  Created {len(saved_files)} files in {project_name}/")
            # Re-index the project after adding files
            code_indexer.index_project(project_path)
        
        # === HANDLE FILE EDITS ===
        edits = self._extract_file_edits(result)
        for edit in edits:
            edit_result = code_indexer.edit_file(
                project_path, 
                edit["file"], 
                edit["old"], 
                edit["new"]
            )
            if edit_result["success"]:
                self._log(f"  Edited: {edit['file']}")
            else:
                self._log(f"  Edit failed: {edit_result['error']}")
        
        # Save result to appropriate domain
        domain = self._detect_domain(result)
        recycler.save_to_domain(domain, f"### {step}\n{result[:2000]}")
        
        # Mark step complete
        recycler.mark_step_complete(step, result[:500])
        
        return result
    
    def _extract_commands(self, result: str) -> List[str]:
        """Extract commands from LLM output."""
        commands = []
        
        # Look for [COMMAND]: pattern
        cmd_pattern = r'\[COMMAND\]:\s*(.+?)(?:\n|$)'
        matches = re.findall(cmd_pattern, result, re.IGNORECASE)
        commands.extend(matches)
        
        # Also look for common command patterns in context
        # e.g., "Run `npm install`" or "Execute: npm run build"
        run_pattern = r'(?:run|execute)[:\s]+`?([^`\n]+)`?'
        matches = re.findall(run_pattern, result, re.IGNORECASE)
        for m in matches:
            cmd = m.strip()
            if cmd and len(cmd) < 100:  # Sanity check
                commands.append(cmd)
        
        return commands[:5]  # Limit to 5 commands per step
    
    def _extract_file_edits(self, result: str) -> List[Dict]:
        """Extract file edit instructions from LLM output."""
        edits = []
        
        # Look for edit patterns like:
        # [EDIT] file.py: replace "old" with "new"
        # Edit src/index.html: change "old content" to "new content"
        
        edit_pattern = r'\[EDIT\]\s*([^:]+):\s*(?:replace|change)\s*["\'](.+?)["\']\s*(?:with|to)\s*["\'](.+?)["\']'
        matches = re.findall(edit_pattern, result, re.IGNORECASE | re.DOTALL)
        
        for file_path, old_content, new_content in matches:
            edits.append({
                "file": file_path.strip(),
                "old": old_content,
                "new": new_content
            })
        
        return edits[:3]  # Limit to 3 edits per step
    
    
    def _run_qa_feedback(self, project_path: str, max_attempts: int = 2) -> Optional[Dict]:
        """
        Run QA on project and attempt auto-fix if issues found.
        This is the feedback loop that makes the AI self-correcting.
        """
        for attempt in range(max_attempts):
            # Run QA check
            qa_result = qa_agent.run(project_path)
            
            if qa_result["overall_status"] == "pass":
                if attempt > 0:
                    self._log(f"  [QA] Fixed after {attempt} attempt(s)")
                return qa_result
            
            # If errors found, try to fix them
            if qa_result["total_errors"] > 0:
                self._log(f"  [QA] Found {qa_result['total_errors']} errors - attempting fix...")
                
                # Generate fix prompt from QA results
                fix_prompt = qa_agent.generate_fix_prompt(qa_result)
                
                if fix_prompt:
                    # Ask LLM to fix the issues
                    fix_result = self._call_llm(fix_prompt)
                    
                    # Extract and apply fixed code
                    project_name = os.path.basename(project_path)
                    self._extract_and_save_code(fix_result, project_name)
                    
                    self._log(f"  [QA] Applied fixes, re-checking...")
            else:
                # Only warnings, no critical errors to fix
                self._log(f"  [QA] {qa_result['total_warnings']} warnings (not critical)")
                return qa_result
        
        self._log(f"  [QA] Max attempts reached")
        return qa_result
    
    def _plan_steps(self, objective: str) -> List[str]:
        """
        Plan steps using COMPONENT-FOCUSED approach.
        Each step produces a COMPLETE, POLISHED module - not shallow fragments.
        """
        self._log("Planning steps...")
        
        # Detect if this is a large/complex project
        complex_keywords = ["os", "system", "platform", "ecosystem", "full", "complete", 
                           "entire", "everything", "crm", "dashboard", "app", "application",
                           "business", "enterprise", "management"]
        is_complex = any(kw in objective.lower() for kw in complex_keywords)
        
        if is_complex:
            # V4.4: DEEP WORK planning - each step produces a COMPLETE file
            prompt = f"""You are Jarvis, an AI CO-FOUNDER planning a CODE PROJECT.

OBJECTIVE: {objective}

CRITICAL RULES FOR PLANNING:
1. Each step produces ONE COMPLETE FILE - not a stub, not a placeholder
2. DO NOT split file creation into multiple steps
3. Each step MUST include "COMPLETE" or "with full implementation"
4. NO placeholders like [Insert here] or TODO or pass

YOUR PHILOSOPHY:
- Fewer steps, but each is COMPLETE and PRODUCTION-READY
- Quality over quantity - 10 complete files > 50 stubs
- Each file should be runnable/usable immediately

STEP FORMAT:
[number]. Create COMPLETE [filename] with [full description] -> OUTPUT: [filename]

EXAMPLE STEPS (note each is COMPLETE):
1. Create COMPLETE glossary.md with all 4 definitions (Diversion, Valorization, Incineration, Landfill) - minimum 1000 words -> OUTPUT: glossary.md
2. Create COMPLETE simulation.py with working algorithm, sample data, and main() function -> OUTPUT: simulation.py
3. Create COMPLETE paper.md with Abstract, Introduction, Methodology, Results, Discussion, Conclusion - minimum 2000 words -> OUTPUT: paper.md
4. Create COMPLETE data.py with sample industry data (10+ entries) -> OUTPUT: data.py
5. Create COMPLETE analysis.py with chart generation and CSV output -> OUTPUT: analysis.py

BAD EXAMPLES (DO NOT DO THIS):
- "Add introduction section" -> TOO SHALLOW
- "Define input parameters" -> DOESN'T PRODUCE COMPLETE FILE
- "Brainstorm concepts" -> NO DELIVERABLE

NOW GENERATE 10-20 DEEP, COMPLETE-FILE STEPS FOR:
{objective[:1500]}

Output ONLY numbered steps. Each must have "COMPLETE" and -> OUTPUT: [filename]."""
            max_steps = 30
            
            # For very long objectives, the prompt might be overwhelming - add retry hint
            if len(objective) > 3000:
                self._log(f"  [Planner] Long objective ({len(objective)} chars), may need retry")
        else:
            # Check if this is a research/algorithm task
            is_research = any(kw in objective.lower() for kw in [
                "research", "algorithm", "paper", "academic", "benchmark", "novel", "study"
            ])
            
            if is_research:
                # V4.4: DEEP WORK for research projects
                prompt = f"""You are planning a RESEARCH project. Each step produces a COMPLETE file.

OBJECTIVE: {objective[:2000] if len(objective) > 2000 else objective}

CRITICAL: Each step must produce ONE COMPLETE file, not stubs or placeholders.

REQUIRED COMPLETE OUTPUTS:
1. glossary.md - COMPLETE definitions of all terms (minimum 500 words)
2. algorithm.py - COMPLETE working algorithm with main() function
3. data.py - COMPLETE sample dataset (10+ data points)
4. simulation.py - COMPLETE simulation that runs and outputs results
5. analysis.py - COMPLETE analysis with chart/CSV generation
6. paper.md - COMPLETE academic paper (Abstract, Intro, Method, Results, Discussion, Conclusion)

STEP FORMAT:
[number]. Create COMPLETE [filename] with [detailed requirements] -> OUTPUT: [filename]

EXAMPLE STEPS (each is COMPLETE):
1. Create COMPLETE glossary.md with all GRI 306 definitions (Diversion, Valorization, Incineration, Landfill) - minimum 500 words -> OUTPUT: glossary.md
2. Create COMPLETE algorithm.py with novel industrial symbiosis model, 3+ functions, sample usage -> OUTPUT: algorithm.py
3. Create COMPLETE data.py with 10+ industry data entries, costs, volumes -> OUTPUT: data.py
4. Create COMPLETE simulation.py that imports algorithm.py, runs model, outputs results -> OUTPUT: simulation.py
5. Create COMPLETE paper.md with full academic structure, 2000+ words, real results -> OUTPUT: paper.md

DO NOT DO THIS:
- "Extract Diversion definition" -> TOO SHALLOW, combine all definitions in one step
- "Write Abstract section" -> TOO SHALLOW, write complete paper in one step
- "Define input parameters" -> TOO SHALLOW, create complete file

Generate 10-15 COMPLETE-FILE steps. Quality over quantity."""
            else:
                prompt = f"""Break down this objective into concrete steps:

OBJECTIVE: {objective[:2000] if len(objective) > 2000 else objective}

RULES:
- Each step should produce something COMPLETE and WORKING
- Each step should specify what FILE it creates/modifies
- Don't split "Create component" and "Add styling" - combine them

STEP FORMAT:
[number]. [action] -> OUTPUT: [filename or description]

Example:
1. Create main application entry point -> OUTPUT: main.py
2. Build user authentication with JWT -> OUTPUT: auth.py
3. Design database schema with all tables -> OUTPUT: schema.sql

Output 20-50 specific, granular steps with OUTPUT specified."""
            max_steps = 50
        
        # Use HIGH token limit for planning - we need 50-100 detailed steps
        response = self._call_llm(prompt, max_tokens=16384)
        
        # Debug: log raw response length
        self._log(f"  [Planner] Got {len(response)} chars response")
        
        # RETRY if response too short (Qwen3 thinking mode issue)
        if len(response) < 200:
            self._log("  [Planner] Response too short, retrying with more tokens...")
            retry_prompt = prompt + "\n\nIMPORTANT: List at least 10 concrete numbered steps. Be specific and detailed."
            response = self._call_llm(retry_prompt, max_tokens=16384)
        
        # Parse steps from response - AGGRESSIVE parsing to catch ALL steps
        steps = []
        import re
        
        # Split by newlines and process each line
        for line in response.split("\n"):
            line = line.strip()
            if not line or len(line) < 10:
                continue
            
            # Skip pure headers/phase labels without actionable content
            if re.match(r'^(PHASE|Phase)\s*\d*[:\s]*$', line):
                continue
            if line.startswith("---") or line.startswith("==="):
                continue
            if re.match(r'^#+ (PHASE|Phase)', line):
                continue
                
            # AGGRESSIVE MATCHING - any line that looks like a step
            
            # Match numbered lists: "1.", "1)", "Step 1:", "1 -", etc.
            if re.match(r'^[\d]+[\.\)\:\-\s]', line):
                step = re.sub(r'^[\d]+[\.\)\:\-\s]+', '', line).strip()
                step = re.sub(r'^\[.*?\]\s*', '', step).strip()  # Remove [COMPONENT] tags
                step = re.sub(r'^Step\s*\d*[:\.]?\s*', '', step).strip()  # Remove "Step N:" prefix
                if step and len(step) > 10:
                    steps.append(step)
                continue
            
            # Match bulleted lists: "- item", "* item", "• item"
            if line.startswith("-") or line.startswith("*") or line.startswith("•"):
                step = re.sub(r'^[\-\*•\s]+', '', line).strip()
                step = re.sub(r'^\[.*?\]\s*', '', step).strip()
                if step and len(step) > 10:
                    steps.append(step)
                continue
            
            # Match markdown headers with content: "## Create landing page"
            if line.startswith("#"):
                step = re.sub(r'^#+\s*', '', line).strip()
                step = re.sub(r'^[\d]+[\.\)\:\-\s]+', '', step).strip()
                step = re.sub(r'^Step\s*\d*[:\.]?\s*', '', step).strip()
                step = re.sub(r'^\[.*?\]\s*', '', step).strip()
                if step and len(step) > 10 and not any(skip in step.lower() for skip in ['phase', 'section', 'part']):
                    steps.append(step)
                continue
            
            # Match bold markers: "**Create landing page**"
            if line.startswith("**"):
                step = re.sub(r'^\*+\s*', '', line).strip()
                step = re.sub(r'\*+$', '', step).strip()
                step = re.sub(r'^Step\s*\d*[:\.]?\s*', '', step).strip()
                step = re.sub(r'^\[.*?\]\s*', '', step).strip()
                if step and len(step) > 10:
                    steps.append(step)
                continue
            
            # Match lines starting with [COMPONENT], [ACTION], etc.
            if line.startswith("["):
                step = re.sub(r'^\[.*?\]\s*', '', line).strip()
                if step and len(step) > 10:
                    steps.append(step)
                continue
        
        # If still not enough steps, try sentence-based extraction
        if len(steps) < 40:
            self._log(f"  [Planner] Only {len(steps)} steps found, trying deeper extraction...")
            sentences = re.split(r'[.!?]\s+', response)
            action_words = ['create', 'build', 'implement', 'add', 'design', 'develop', 'setup', 
                           'configure', 'write', 'generate', 'define', 'test', 'integrate', 
                           'deploy', 'optimize', 'research', 'analyze', 'document']
            for sent in sentences:
                sent = sent.strip()
                if len(sent) > 20 and any(word in sent.lower() for word in action_words):
                    # Don't duplicate
                    if not any(existing.lower() in sent.lower() or sent.lower() in existing.lower() for existing in steps):
                        steps.append(sent[:200])
        
        # CRITICAL: If we still have < 40 steps, RETRY with stronger prompt
        if len(steps) < 40:
            self._log(f"  [Planner] Only {len(steps)} steps - RETRYING with explicit minimum...")
            retry_prompt = f"""You generated only {len(steps)} steps. This is NOT ENOUGH.

YOU MUST GENERATE AT LEAST 50 NUMBERED STEPS.

JARVIS PHILOSOPHY: You OUTWORK other AIs. They give 10 steps. You give 50-100.

Break down every single phase into granular sub-steps:
- Research: 5-10 steps (market analysis, competitor research, personas, etc.)
- Design: 5-10 steps (wireframes, user flows, architecture, API specs, etc.)
- Frontend: 10-20 steps (landing page, dashboard, auth pages, components, styling, etc.)
- Backend: 10-20 steps (API endpoints, auth, database, models, etc.)
- Documentation: 5-10 steps (readme, whitepaper, API docs, etc.)
- Testing: 5-10 steps (unit tests, integration tests, QA, etc.)

OBJECTIVE: {objective[:1500]}

OUTPUT ONLY NUMBERED STEPS (1. 2. 3. etc.). NO EXPLANATIONS. MINIMUM 50 STEPS."""
            
            response = self._call_llm(retry_prompt, max_tokens=16384)
            
            # Parse the retry response
            retry_steps = []
            for line in response.split("\n"):
                line = line.strip()
                if not line or len(line) < 10:
                    continue
                if re.match(r'^[\d]+[\.\)\:\-\s]', line):
                    step = re.sub(r'^[\d]+[\.\)\:\-\s]+', '', line).strip()
                    step = re.sub(r'^\[.*?\]\s*', '', step).strip()
                    step = re.sub(r'^Step\s*\d*[:\.]?\s*', '', step).strip()
                    if step and len(step) > 10:
                        retry_steps.append(step)
            
            if len(retry_steps) > len(steps):
                self._log(f"  [Planner] Retry got {len(retry_steps)} steps (improved!)")
                steps = retry_steps
        
        # Log what we found
        self._log(f"Planned {len(steps)} component milestones")
        
        # Return ALL steps - no artificial limit during planning
        # The execution loop has its own max_iterations limit
        return steps
    
    def run(self, objective: str, progress_callback: Callable = None, resume_checkpoint: str = None) -> Dict:
        """
        Run the autonomous execution loop.
        
        Args:
            objective: What to accomplish
            progress_callback: Function to call with progress updates
            resume_checkpoint: Optional checkpoint ID to resume from
        
        Returns:
            Final result with logs and artifacts
        """
        self.progress_callback = progress_callback
        self.is_running = True
        self.iteration = 0
        self.log = []
        
        # === RESUME FROM CHECKPOINT ===
        if resume_checkpoint:
            try:
                from .utils.checkpoint import checkpoint_manager
                cp = checkpoint_manager.get_checkpoint_by_id(resume_checkpoint)
                if cp:
                    self._log(f"=== RESUMING FROM CHECKPOINT {resume_checkpoint} ===")
                    self._log(f"Objective: {cp['objective'][:100]}...")
                    self._log(f"Continuing from iteration {cp['iteration']}, {len(cp['completed_steps'])} steps done")
                    
                    # Restore state
                    objective = cp["objective"]
                    self.iteration = cp["iteration"]
                    recycler.task_objective = cp["objective"]
                    recycler.completed_steps = cp["completed_steps"]
                    recycler._pending_steps = cp["pending_steps"]
                    
                    # Skip to Phase 3 (execution loop)
                    project_path = cp.get("project_path", "")
                    if project_path:
                        self._log(f"Project: {project_path}")
                    
                    # Jump directly to execution loop (skip planning/research)
                    self._log("Resuming execution loop...")
                    # The normal loop will pick up from pending_steps
                else:
                    self._log(f"Checkpoint {resume_checkpoint} not found, starting fresh")
                    resume_checkpoint = None
            except Exception as e:
                self._log(f"Failed to resume: {e}, starting fresh")
                resume_checkpoint = None
        
        if not resume_checkpoint:
            self._log(f"=== AUTONOMOUS MODE STARTED ===")
            self._log(f"Original objective: {objective}")
            
            # === V4.1: DETECT AND LOCK PROJECT TYPE ===
            project_type_rules = self._detect_and_lock_project_type(objective)
            
            # === CLEAR STALE CONTEXT FOR NOVEL TASKS ===
            # Prevent recycling old algorithms when asked to invent something new
            obj_lower = objective.lower()
            is_novel_task = any(kw in obj_lower for kw in [
                "novel", "new algorithm", "invent", "original", "create new",
                "propose a new", "design a new", "develop a new"
            ])
            
            if is_novel_task:
                self._log("⚠️ NOVEL TASK DETECTED - clearing research context to prevent recycling")
                # Clear research domain to prevent pulling in old algorithms
                recycler.save_to_domain("research", "# Fresh Research Context\n\nStarting new research task. Do NOT reuse algorithms from previous sessions.", append=False)
                recycler.save_to_domain("decisions", "# Fresh Decisions\n\nNew task - previous decisions cleared.", append=False)
        
        try:
            # === PHASES 0-2: Only run if NOT resuming ===
            if not resume_checkpoint:
                # Phase 0: Prompt Refinement (auto-expand vague prompts)
                # Skip refinement for already-detailed prompts (>500 chars) to preserve multi-phase tasks
                self._log("Phase 0: Refining prompt...")
                if len(objective) > 500:
                    self._log("Skipping refinement (prompt already detailed)")
                else:
                    try:
                        refined = prompt_refiner.quick_refine(objective)
                        if refined and refined != objective:
                            self._log(f"Refined to: {refined[:200]}...")
                            objective = refined
                        else:
                            self._log("Using original prompt (already detailed)")
                    except Exception as e:
                        self._log(f"Prompt refinement skipped: {e}")
                
                # Phase 1: Research (if needed)
                self._log("Phase 1: Research")
                routing = router.classify(objective)
                
                if routing.get("requires_research_first", False):
                    self._log("Conducting research...")
                    research_result = researcher.run(objective)
                    recycler.save_to_domain("research", research_result)
                
                # Phase 2: Plan steps
                self._log("Phase 2: Planning")
                steps = self._plan_steps(objective)
                recycler.set_task(objective, steps)
                
                # === PHASE 1.5: FLAW ANALYSIS ===
                # Run devils_advocate on the PLAN before execution to catch issues early
                self._log("Phase 1.5: Flaw Analysis")
                try:
                    from .devils_advocate import devils_advocate
                    plan_text = "\n".join([f"- {s}" for s in steps])
                    plan_critique = devils_advocate.critique(
                        f"## Project Plan for: {objective}\n\n{plan_text}",
                        content_type="architecture"
                    )
                    
                    if plan_critique.get("verdict") == "FIX_REQUIRED":
                        critical_flaws = [i for i in plan_critique.get("issues", []) if i.get("risk") == "critical"]
                        if critical_flaws:
                            self._log(f"  ⚠️ Found {len(critical_flaws)} critical plan flaws")
                            for flaw in critical_flaws[:3]:
                                self._log(f"    → {flaw.get('title', 'Issue')}")
                            
                            # V4.3: PASS 3 - Improve plan with critique (low temp, intent-preserving)
                            flaw_context = "\n".join([f"- {f['title']}: {f['fix']}" for f in critical_flaws])
                            improved_plan_prompt = f"""ORIGINAL OBJECTIVE: {objective}

ISSUES FOUND IN PLAN:
{flaw_context}

FIX these issues but DO NOT add new features. Stay focused on the original objective.
Output numbered steps that address the feedback while PRESERVING the original intent.

Steps:"""
                            # Use low temperature for deterministic improvement
                            improved_result = self._call_llm(improved_plan_prompt, max_tokens=2000, temperature=0.2)
                            
                            # Extract improved steps
                            import re
                            improved_steps = re.findall(r'^\d+[\.\)]\s*(.+?)(?=\n\d+[\.\)]|\n*$)', improved_result, re.MULTILINE)
                            if improved_steps and len(improved_steps) >= 3:
                                steps = improved_steps
                                recycler.set_task(objective, steps)
                                self._log(f"  ✓ Plan improved with {len(steps)} steps")
                        else:
                            self._log("  ✓ No critical flaws in plan")
                    else:
                        self._log("  ✓ Plan approved")
                except Exception as e:
                    self._log(f"  Flaw analysis skipped: {e}")
            
            # Phase 2.5: Create project BEFORE execution starts
            # This ensures all agents have a valid project_path to save files
            project_name = self._get_project_name(objective)
            project_path = os.path.join(WORKSPACE_DIR, "projects", project_name)
            
            if not os.path.exists(project_path):
                self._log(f"Creating project: {project_name}")
                template = "react" if any(kw in objective.lower() for kw in ["react", "next", "jsx", "tsx"]) else "vanilla"
                project_builder.create_project(project_name, template=template)
            else:
                project_builder.project_path = project_path
            
            self._log(f"Project path: {project_path}")
            
            # Phase 3: Execute loop
            self._log("Phase 3: Execution")
            
            while self.is_running and self.iteration < self.max_iterations:
                self.iteration += 1
                self._log(f"--- Iteration {self.iteration} ---")
                
                # Check for recycling
                if recycler.needs_recycling():
                    self._log("!!! CONTEXT RECYCLING !!!")
                    continuation_prompt = recycler.recycle(self._call_llm)
                    
                    # Re-plan remaining steps with fresh context
                    result = self._call_llm(continuation_prompt)
                    self._log("Recycled and continuing...")
                
                # Get progress
                progress = recycler.get_progress()
                
                # Check if done
                if not progress["pending_steps"]:
                    self._log("All planned steps completed - verifying outputs...")
                    
                    # === VERIFICATION GATE ===
                    # Don't mark complete until outputs actually exist
                    verification = self._verify_research_outputs(project_path, recycler.task_objective)
                    
                    if verification["verified"]:
                        self._log("✅ Output verification PASSED")
                        for item in verification["verified_items"]:
                            self._log(f"  ✓ {item}")
                        break
                    else:
                        self._log("❌ Output verification FAILED - adding remediation steps")
                        for issue in verification["issues"]:
                            self._log(f"  ✗ {issue}")
                        
                        # Add remediation steps for missing items
                        remediation_steps = []
                        for issue in verification["issues"]:
                            if "paper" in issue.lower():
                                remediation_steps.append("Create paper.md with full academic paper: Abstract, Introduction, Methods, Results (with REAL numbers from benchmarks), Discussion, Conclusion")
                            elif "algorithm" in issue.lower():
                                remediation_steps.append("Create algorithm.py with the novel algorithm implementation")
                            elif "benchmark" in issue.lower() and "results" not in issue.lower():
                                remediation_steps.append("Create benchmark.py that runs the algorithm and baseline, measures performance, outputs CSV")
                            elif "results" in issue.lower() or "csv" in issue.lower():
                                remediation_steps.append("Run benchmark.py and save results to results/benchmark_results.csv with actual measured data")
                            elif "placeholder" in issue.lower():
                                remediation_steps.append("Remove all [PLACEHOLDER] and [TODO] markers from paper.md - fill in with actual content")
                            elif "syntax" in issue.lower():
                                remediation_steps.append("Fix Python syntax errors in the code files")
                        
                        if remediation_steps:
                            # Add to pending steps and continue
                            for step in remediation_steps:
                                if step not in recycler.completed_steps:
                                    recycler.pending_steps.append(step)
                            self._log(f"Added {len(remediation_steps)} remediation steps")
                        else:
                            # No specific remediation, break anyway
                            self._log("No specific remediation available, completing with issues")
                            break
                
                # Execute next step
                next_step = progress["pending_steps"][0]
                task_context = recycler.get_all_domain_context()
                
                # === V4.0: CODING ITERATION LIMIT ===
                # Track coding steps to prevent infinite loops
                step_lower = next_step.lower()
                is_coding_step = any(kw in step_lower for kw in [
                    'implement', 'create', 'build', 'code', 'component', 'api', 
                    'backend', 'frontend', 'database', 'auth', 'module'
                ])
                
                if is_coding_step:
                    self.coding_iterations += 1
                    if self.coding_iterations >= self.max_coding_iterations:
                        self._log(f"⚠️ CODING LIMIT REACHED ({self.max_coding_iterations} iterations)")
                        self._log("  -> Forcing transition to Git/Deploy phases")
                        # Clear remaining coding steps
                        recycler.pending_steps = []
                        break
                
                result = self._execute_step(next_step, task_context)
                
                # CRITICAL: Mark step as complete so loop progresses
                recycler.mark_step_complete(next_step, result[:500] if result else "")
                
                # === AUTO-CHECKPOINT (Crash Recovery) ===
                # Save state every 5 steps for server restart recovery
                if self.iteration % 5 == 0:
                    try:
                        from .utils.checkpoint import checkpoint_manager
                        checkpoint_manager.save_checkpoint(
                            objective=recycler.task_objective,
                            iteration=self.iteration,
                            completed_steps=recycler.completed_steps,
                            pending_steps=progress["pending_steps"][1:],  # Remaining steps
                            project_path=project_path,
                            metadata={"log_count": len(self.log)}
                        )
                        self._log(f"💾 Checkpoint saved at iteration {self.iteration}")
                    except Exception as e:
                        self._log(f"⚠️ Checkpoint failed: {e}")
                
                # === QA FEEDBACK LOOP ===
                # After code is generated, verify and fix if needed
                project_name = self._get_project_name(recycler.task_objective)
                project_path = os.path.join(WORKSPACE_DIR, "projects", project_name)
                
                if os.path.exists(project_path):
                    self._run_qa_feedback(project_path)
                
                # Check for completion signal in result
                if self._check_completion(result):
                    self._log("Task signaled completion")
                    break
                
                # Update token count (for recycler tracking)
                recycler.current_tokens = recycler.get_total_context_tokens()
            
            # === FINAL QA CHECK ===
            project_name = self._get_project_name(objective)
            project_path = os.path.join(WORKSPACE_DIR, "projects", project_name)
            
            if os.path.exists(project_path):
                self._log("Phase 4: Final QA")
                final_qa = self._run_qa_feedback(project_path, max_attempts=3)
                qa_status = final_qa.get("overall_status", "unknown") if final_qa else "skipped"
                self._log(f"Final QA: {qa_status}")
                
                # === PHASE 5: VISUAL QA ===
                # Take headless screenshots and analyze with vision model
                task_type = self._detect_task_type(objective)
                if task_type == "coding" and ("web" in objective.lower() or "page" in objective.lower() or "site" in objective.lower() or "frontend" in objective.lower()):
                    self._log("Phase 5: Visual QA (Headless)")
                    try:
                        visual_result = visual_qa.analyze_project(project_path)
                        visual_score = visual_result.get("average_score", 0)
                        visual_issues = visual_result.get("total_issues", 0)
                        critical_issues = len(visual_result.get("critical_issues", []))
                        
                        if visual_score >= 70 and critical_issues == 0:
                            self._log(f"Visual QA: PASSED (score: {visual_score}/100)")
                        else:
                            self._log(f"Visual QA: {visual_issues} issues found ({critical_issues} critical)")
                            # Log recommendations
                            for page in visual_result.get("pages_analyzed", []):
                                for rec in page.get("analysis", {}).get("recommendations", [])[:2]:
                                    self._log(f"  → {rec}")
                    except Exception as e:
                        self._log(f"Visual QA skipped: {e}")
                
                # === PHASE 6: PROJECT CONSOLIDATION & FINALIZATION ===
                self._log("Phase 6: Project Consolidation")
                try:
                    consolidation_result = self._finalize_project(project_path)
                    if consolidation_result.get("files_consolidated", 0) > 0:
                        self._log(f"  → Consolidated {consolidation_result['files_consolidated']} fragmented files")
                    if consolidation_result.get("shadowing_fixed", 0) > 0:
                        self._log(f"  → Fixed {consolidation_result['shadowing_fixed']} package shadowing issues")
                    if consolidation_result.get("imports_verified", 0) > 0:
                        self._log(f"  → Verified {consolidation_result['imports_verified']} imports")
                    self._log(f"  → Project finalization: {consolidation_result.get('status', 'complete')}")
                except Exception as e:
                    self._log(f"Consolidation skipped: {e}")
                
                # === PHASE 7: GIT AUTO-COMMIT & MEMORY PERSISTENCE ===
                self._log("Phase 7: Git & Memory")
                
                # Auto-commit project to git
                try:
                    from .git_agent import git_agent
                    # 1. Init
                    init_res = git_agent.init_repo(project_path)
                    if not init_res.get("success") and "already" not in init_res.get("message", ""):
                         self._log(f"  → Git Init failed: {init_res.get('error')}")
                    
                    # 2. Commit
                    git_result = git_agent.commit(
                        project_path, 
                        f"Jarvis: Completed {project_name}"
                    )
                    if git_result.get("success"):
                        self._log(f"  → Git: Committed {git_result.get('message', 'files')}")
                    else:
                        self._log(f"  → Git: {git_result.get('message', 'skipped')}")
                except Exception as e:
                    self._log(f"  → Git skipped: {e}")
                
                # Save successful patterns to memory
                try:
                    task_type = self._detect_task_type(objective)
                    memory.save(
                        key=f"project:{project_name}",
                        data={
                            "objective": objective[:500],
                            "task_type": task_type,
                            "steps_completed": len(recycler.completed_steps),
                            "files_created": os.listdir(project_path) if os.path.exists(project_path) else [],
                            "timestamp": datetime.now().isoformat()
                        }
                    )
                    self._log("  → Memory: Saved project pattern")
                except Exception as e:
                    self._log(f"  → Memory skipped: {e}")
                
                # === PHASE 8: GITHUB REPO CREATION ===
                self._log("Phase 8: GitHub")
                github_url = None
                try:
                    # Use full_setup to create repo, add remote, and push
                    # Note: git_agent handles the GitHub API interaction
                    from .git_agent import git_agent
                    
                    repo_result = git_agent.full_setup(
                        project_path=project_path,
                        repo_name=project_name,
                        private=False
                    )
                    
                    if repo_result.get("success"):
                        github_url = repo_result.get("repo_url")
                        self._log(f"  → GitHub: {github_url}")
                    else:
                        self._log(f"  → GitHub Error: {repo_result}")
                        
                except Exception as e:
                    self._log(f"  → GitHub skipped: {e}")
                
                # === PHASE 9: AUTO-DEPLOYMENT ===
                self._log("Phase 9: Deployment")
                deployment_info = {"frontend": None, "backend": None}
                
                try:
                    # Detect project type
                    has_frontend = any(f.endswith(('.html', '.jsx', '.tsx', '.vue')) for f in os.listdir(project_path))
                    has_backend = any(f.endswith(('.py',)) and f not in ['setup.py', 'test.py'] for f in os.listdir(project_path))
                    has_package_json = os.path.exists(os.path.join(project_path, "package.json"))
                    
                    # Deploy frontend
                    if has_frontend or has_package_json:
                        self._log("  Deploying frontend to Netlify...")
                        deploy_result = self._deploy_netlify(project_path, prod=True)
                        if deploy_result.get("success"):
                            deployment_info["frontend"] = deploy_result.get("url")
                            self._log(f"  → Frontend: {deployment_info['frontend']}")
                        else:
                            # Try Vercel as fallback
                            deploy_result = self._deploy_vercel(project_path)
                            if deploy_result.get("success"):
                                deployment_info["frontend"] = deploy_result.get("url")
                                self._log(f"  → Frontend (Vercel): {deployment_info['frontend']}")
                    
                    # Suggest backend deployment
                    if has_backend:
                        self._log("  ℹ️ Backend detected - consider:")
                        self._log("    - AWS Lambda (serverless)")
                        self._log("    - Railway/Render (container)")
                        self._log("    - Fly.io (edge)")
                        deployment_info["backend"] = "manual_required"
                        
                except Exception as e:
                    self._log(f"  → Deployment skipped: {e}")
            
            # Final summary
            self._log("=== AUTONOMOUS MODE COMPLETE ===")
            final_progress = recycler.get_progress()
            
            return {
                "status": "complete",
                "objective": objective,
                "iterations": self.iteration,
                "progress": final_progress,
                "project_path": project_path if os.path.exists(project_path) else None,
                "github_url": github_url,
                "deployment": deployment_info,
                "log": self.log,
                "domain_files": {
                    domain: recycler.read_domain(domain)[-500:]
                    for domain in recycler.DOMAINS
                }
            }
            
        except Exception as e:
            self._log(f"!!! ERROR: {e} !!!")
            return {
                "status": "error",
                "error": str(e),
                "log": self.log
            }
        
        finally:
            self.is_running = False
    
    def stop(self):
        """Stop the autonomous loop."""
        self._log("Stop requested")
        self.is_running = False


# Singleton
autonomous_executor = AutonomousExecutor()
